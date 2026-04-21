"""
將兩本電子書 Markdown 轉換為格式化 Word 檔
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, os

# ── 顏色設定 ──────────────────────────────────────
ACCENT   = RGBColor(0xC4, 0x62, 0x2D)   # 橘棕色
CHARCOAL = RGBColor(0x2C, 0x2C, 0x2C)   # 深灰
MUTED    = RGBColor(0x88, 0x84, 0x7E)   # 灰
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
SAND     = RGBColor(0xF5, 0xF0, 0xE8)   # 米色底

def set_cell_bg(cell, hex_color: str):
    """設定表格儲存格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_horizontal_rule(doc):
    """加入分隔線"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C4622D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(12)
    return p

def setup_doc_styles(doc):
    """設定全域頁面與樣式"""
    section = doc.sections[0]
    section.page_width  = Inches(8.27)   # A4
    section.page_height = Inches(11.69)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # 預設字體
    style = doc.styles['Normal']
    font  = style.font
    font.name = 'Noto Sans TC'
    font.size = Pt(11)
    font.color.rgb = CHARCOAL
    style.paragraph_format.space_after = Pt(6)

def add_cover(doc, title: str, subtitle: str, author: str, price: str):
    """封面頁"""
    # 橘色大標
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    run = p.add_run(title)
    run.font.size  = Pt(26)
    run.font.bold  = True
    run.font.color.rgb = ACCENT

    # 副標
    p2 = doc.add_paragraph(subtitle)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].font.size  = Pt(14)
    p2.runs[0].font.color.rgb = MUTED
    p2.paragraph_format.space_before = Pt(8)

    add_horizontal_rule(doc)

    # 作者
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p3.add_run(author)
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = CHARCOAL

    # 定價
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rp = p4.add_run(price)
    rp.font.size  = Pt(11)
    rp.font.color.rgb = MUTED

    doc.add_page_break()


def parse_and_build(doc, md_path: str):
    """解析 Markdown 並逐行寫入 Word"""
    with open(md_path, encoding='utf-8') as f:
        lines = f.readlines()

    in_code    = False
    code_lines = []
    in_table   = False
    table_rows = []
    skip_cover = True   # 跳過前幾行的封面資訊（已另外用 add_cover 產生）
    cover_end  = 0

    # 找封面結束位置（第一個 --- 之後再一個 ---）
    dash_count = 0
    for i, line in enumerate(lines):
        if line.strip() == '---':
            dash_count += 1
            if dash_count == 2:
                cover_end = i + 1
                break

    lines = lines[cover_end:]   # 跳過封面

    i = 0
    while i < len(lines):
        raw  = lines[i]
        line = raw.rstrip('\n')

        # ── 程式碼區塊 ──────────────────────────────
        if line.startswith('```'):
            if not in_code:
                in_code    = True
                code_lines = []
            else:
                in_code = False
                # 輸出程式碼區塊
                p = doc.add_paragraph()
                p.paragraph_format.left_indent  = Cm(1)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after  = Pt(4)
                # 灰底框
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'),   'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'),  'F0EFED')
                pPr.append(shd)
                # 左邊框線
                pBdr = OxmlElement('w:pBdr')
                left = OxmlElement('w:left')
                left.set(qn('w:val'),   'single')
                left.set(qn('w:sz'),    '12')
                left.set(qn('w:space'), '4')
                left.set(qn('w:color'), 'C4622D')
                pBdr.append(left)
                pPr.append(pBdr)

                code_text = '\n'.join(code_lines)
                run = p.add_run(code_text)
                run.font.name  = 'Courier New'
                run.font.size  = Pt(9.5)
                run.font.color.rgb = RGBColor(0x40, 0x3C, 0x38)
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── 表格 ──────────────────────────────────────
        if line.startswith('|'):
            table_rows.append(line)
            i += 1
            continue
        elif table_rows:
            # 渲染收集到的表格
            # 過濾掉分隔列 |---|---|
            rows = [r for r in table_rows if not re.match(r'^\|[-| :]+\|$', r.strip())]
            if rows:
                parsed = []
                for tr in rows:
                    cells = [c.strip() for c in tr.strip().strip('|').split('|')]
                    parsed.append(cells)

                cols = max(len(r) for r in parsed)
                tbl  = doc.add_table(rows=len(parsed), cols=cols)
                tbl.style = 'Table Grid'
                tbl.autofit = True

                for ri, row_data in enumerate(parsed):
                    for ci, cell_text in enumerate(row_data):
                        if ci >= cols:
                            break
                        cell = tbl.cell(ri, ci)
                        cell.text = ''
                        p_cell = cell.paragraphs[0]
                        p_cell.paragraph_format.space_before = Pt(2)
                        p_cell.paragraph_format.space_after  = Pt(2)

                        # 表頭列
                        if ri == 0:
                            set_cell_bg(cell, 'C4622D')
                            run = p_cell.add_run(cell_text)
                            run.font.bold  = True
                            run.font.color.rgb = WHITE
                            run.font.size  = Pt(10)
                        else:
                            set_cell_bg(cell, 'FDFCFB' if ri % 2 == 0 else 'F5F0E8')
                            run = p_cell.add_run(cell_text)
                            run.font.size  = Pt(10)
                            run.font.color.rgb = CHARCOAL

            doc.add_paragraph()   # 表格後空行
            table_rows = []
            continue

        # ── 水平分隔線 ---  ───────────────────────────
        if line.strip() == '---':
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── 標題 ──────────────────────────────────────
        heading_match = re.match(r'^(#{1,4})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            text  = heading_match.group(2).strip()
            p = doc.add_paragraph()

            if level == 1:
                p.paragraph_format.space_before = Pt(24)
                p.paragraph_format.space_after  = Pt(8)
                run = p.add_run(text)
                run.font.size  = Pt(22)
                run.font.bold  = True
                run.font.color.rgb = ACCENT

            elif level == 2:
                p.paragraph_format.space_before = Pt(20)
                p.paragraph_format.space_after  = Pt(6)
                # 橘色左邊框
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                left = OxmlElement('w:left')
                left.set(qn('w:val'),   'single')
                left.set(qn('w:sz'),    '16')
                left.set(qn('w:space'), '8')
                left.set(qn('w:color'), 'C4622D')
                pBdr.append(left)
                pPr.append(pBdr)
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run(text)
                run.font.size  = Pt(16)
                run.font.bold  = True
                run.font.color.rgb = CHARCOAL

            elif level == 3:
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after  = Pt(4)
                run = p.add_run(text)
                run.font.size  = Pt(13)
                run.font.bold  = True
                run.font.color.rgb = ACCENT

            else:
                p.paragraph_format.space_before = Pt(10)
                run = p.add_run(text)
                run.font.size  = Pt(11.5)
                run.font.bold  = True
                run.font.color.rgb = CHARCOAL

            i += 1
            continue

        # ── 引用區塊 >  ───────────────────────────────
        if line.startswith('>'):
            quote_text = line.lstrip('> ').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Cm(1.2)
            p.paragraph_format.right_indent = Cm(1.2)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(6)
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  'F5F0E8')
            pPr.append(shd)
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'),   'single')
            left.set(qn('w:sz'),    '16')
            left.set(qn('w:space'), '8')
            left.set(qn('w:color'), 'C4622D')
            pBdr.append(left)
            pPr.append(pBdr)
            inline_format(p, quote_text, italic=True, color=CHARCOAL)
            i += 1
            continue

        # ── 無序列表 - / * / ・  ──────────────────────
        bullet_match = re.match(r'^(\s*)[\-\*・]\s+(.*)', line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            text   = bullet_match.group(2)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent  = Cm(0.5 + indent * 0.4)
            p.paragraph_format.space_after  = Pt(3)
            inline_format(p, text)
            i += 1
            continue

        # ── Checkbox 清單 - [ ]  ──────────────────────
        cb_match = re.match(r'^\s*-\s+\[[ x]\]\s+(.*)', line)
        if cb_match:
            text = '□ ' + cb_match.group(1)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_after = Pt(3)
            inline_format(p, text)
            i += 1
            continue

        # ── 空行 ──────────────────────────────────────
        if line.strip() == '':
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # ── 一般段落 ──────────────────────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        inline_format(p, line)
        i += 1

    return doc


def inline_format(p, text: str, italic: bool = False, color: RGBColor = None):
    """處理行內的 **bold**、*italic*、`code` 格式"""
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|___.*?___)')
    parts   = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.color.rgb = color or CHARCOAL
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = p.add_run(part[1:-1])
            run.italic = True
            run.font.color.rgb = color or CHARCOAL
        elif part.startswith('`') and part.endswith('`') and len(part) > 2:
            run = p.add_run(part[1:-1])
            run.font.name  = 'Courier New'
            run.font.size  = Pt(10)
            run.font.color.rgb = ACCENT
        else:
            run = p.add_run(part)
            if italic:
                run.italic = True
            run.font.color.rgb = color or CHARCOAL
        run.font.size = run.font.size or Pt(11)


# ── 主程式 ────────────────────────────────────────
BASE = os.path.dirname(os.path.abspath(__file__))

books = [
    {
        'md':       os.path.join(BASE, '跨產業求職履歷重寫指南.md'),
        'out':      os.path.join(BASE, '跨產業求職履歷重寫指南.docx'),
        'title':    '跨產業求職\n履歷重寫指南',
        'subtitle': '不是範本，是思考框架',
        'price':    'NT$299　版權所有・蒲朝棟 Tim・職涯停看聽',
    },
    {
        'md':       os.path.join(BASE, '面試高分回答框架10則.md'),
        'out':      os.path.join(BASE, '面試高分回答框架10則.docx'),
        'title':    '面試高分回答\n框架 10 則',
        'subtitle': '從自我介紹到薪資談判，每題附框架與範例解析',
        'price':    'NT$199　版權所有・蒲朝棟 Tim・職涯停看聽',
    },
]

author = 'CDA 認證職涯發展師 蒲朝棟 Tim\n職涯停看聽｜LINE：@tzlth'

for book in books:
    doc = Document()
    setup_doc_styles(doc)
    add_cover(doc, book['title'], book['subtitle'], author, book['price'])
    parse_and_build(doc, book['md'])
    doc.save(book['out'])
    print(f"OK: {book['out']}")

print('Done.')
