"""
產生「讓履歷被看見的關鍵句型_免費精華版.docx」
v3 — 依第二次評估報告修正（2026-04-25）
  v2 既有修正（保留）：
    - 移除「7秒」改 ATS 現實說法 / 選題邏輯 / 沒有數字小框
    - NT$24萬估算邏輯 / ATS關聯性 / 護理師範例 / CTA定位說明
  v3 新增修正：
    Fix A：移除「眼動研究顯示」（無來源），改為保守說法
    Fix B：B2B 範例「年營收」→「年化合約金額」（避免歧義）
    Fix C：PAR 例一 NT$30萬 → NT$14萬（統一 NT$2,000/天假設），補 † 腳注
    Fix D：PAR 例二 客戶留存率補基準（從 72% 提升至 90%）
    Fix E：結語 AI 診斷 CTA 視覺提升（淡藍底色框，轉換障礙最低入口）
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── 頁面設定 ──────────────────────────────────────
section = doc.sections[0]
section.page_width    = Cm(21)
section.page_height   = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── 顏色常數 ──────────────────────────────────────
DARK   = RGBColor(0x1A, 0x1A, 0x2E)
ACCENT = RGBColor(0x0F, 0x72, 0xB1)
GRAY   = RGBColor(0x60, 0x60, 0x60)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
RED    = RGBColor(0xC0, 0x39, 0x2B)
GREEN  = RGBColor(0x1A, 0x7A, 0x4A)
ORANGE = RGBColor(0xE6, 0x7E, 0x22)

# ── 工具函式 ──────────────────────────────────────
def set_font(run, size, bold=False, color=None, italic=False):
    run.font.name = "微軟正黑體"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    rPr.insert(0, rFonts)

def para(text, size=11, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
         space_before=0, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(size * 1.6)
    run = p.add_run(text)
    set_font(run, size, bold, color or DARK, italic)
    return p

def heading1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '0F72B1')
    pPr.append(shd)
    run = p.add_run(f'  {text}')
    set_font(run, 16, bold=True, color=WHITE)
    p.paragraph_format.line_spacing = Pt(28)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, 13, bold=True, color=ACCENT)
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, 11, bold=True, color=DARK)
    return p

def formula_box(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EBF4FB')
    pPr.append(shd)
    run = p.add_run(text)
    set_font(run, 11, bold=True, color=ACCENT)
    p.paragraph_format.line_spacing = Pt(20)
    return p

def bullet(text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size, color=DARK)
    return p

def shaded_para(text, fill_hex, size=10, bold=False, color=None,
                left_indent=0.5, space_before=0, space_after=3):
    """帶底色的段落（用於小框內容行）"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)
    p.paragraph_format.left_indent  = Cm(left_indent)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(size * 1.65)
    run = p.add_run(text)
    set_font(run, size, bold=bold, color=color or DARK)
    return p

def add_tip_box(title, lines, fill='FFF3CD', title_color=None):
    """橘黃色提示小框"""
    tc = title_color or ORANGE
    shaded_para(f'💡 {title}', fill, size=10, bold=True, color=tc,
                space_before=8, space_after=2)
    for line in lines:
        shaded_para(line, fill, size=9.5, color=DARK, space_before=0, space_after=3)

def add_footnote(text):
    """表格下方的小字估算說明"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.3)
    run = p.add_run(text)
    set_font(run, 8.5, color=GRAY, italic=True)

def add_comparison_table(rows_data):
    table = doc.add_table(rows=1 + len(rows_data), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for cell, txt, fill in zip(hdr, ['#', '改前 ❌', '改後 ✅'],
                                ['2C3E50', 'C0392B', '1A7A4A']):
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(txt)
        set_font(run, 10, bold=True, color=WHITE)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        tcPr.append(shd)
    widths = [Cm(0.8), Cm(5.5), Cm(9.5)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(widths[i].pt * 20)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
    for i, (num, before, after) in enumerate(rows_data):
        row = table.rows[i + 1]
        for cell, txt, bold, color in zip(
                row.cells, [str(num), before, after],
                [True, False, False], [DARK, RED, GREEN]):
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(txt)
            set_font(run, 9.5, bold=bold, color=color)
            cell.paragraphs[0].paragraph_format.line_spacing = Pt(15)

def add_practice_box(lines):
    para('', size=4, space_before=0, space_after=2)
    shaded_para('✏️ 練習格式', 'F8F9FA', size=10, bold=True, color=ACCENT,
                space_before=6, space_after=2)
    for line in lines:
        shaded_para(line, 'F8F9FA', size=10, color=DARK,
                    space_before=0, space_after=3)

def add_no_number_box():
    """修正3：沒有數字怎麼辦 小框"""
    add_tip_box(
        '沒有數字怎麼辦？',
        [
            '如果無法取得精確數字，可以用以下方式替代：',
            '・規模（例：服務 50+ 位客戶、管理 3 人小組）',
            '・頻率（例：每週主持 3 場跨部門會議）',
            '・相對比較（例：從每週 2 次縮短至每週 1 次以下）',
            '・定性影響（例：首次建立 XX 機制，此前無標準流程）',
            '',
            '有方向的估算，比空白好。只要推算邏輯合理，估算值也可以寫進履歷。',
        ],
        fill='FFF8E1',
        title_color=ORANGE,
    )

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    p._p.get_or_add_pPr().append(pBdr)

# ══════════════════════════════════════════════════
#  封面頁
# ══════════════════════════════════════════════════
p_cover = doc.add_paragraph()
p_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_cover.paragraph_format.space_before = Pt(60)
p_cover.paragraph_format.space_after  = Pt(8)
run = p_cover.add_run('讓履歷被看見的關鍵句型')
set_font(run, 28, bold=True, color=ACCENT)

para('免費精華版', size=14, align=WD_ALIGN_PARAGRAPH.CENTER,
     color=GRAY, space_before=0, space_after=4)
para('2 個最高頻使用的改寫公式，附改寫前後對照',
     size=11, align=WD_ALIGN_PARAGRAPH.CENTER,
     color=GRAY, space_before=0, space_after=40)
para('蒲朝棟 Tim', size=12, bold=True,
     align=WD_ALIGN_PARAGRAPH.CENTER, color=DARK, space_before=0, space_after=4)
para('CDA 認證職涯顧問・104 職涯引導師',
     size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, space_before=0, space_after=60)

doc.add_page_break()

# ══════════════════════════════════════════════════
#  給你的一段話（修正 1 + 修正 2）
# ══════════════════════════════════════════════════
heading1('給你的一段話')

# 修正 1：移除「7秒」，改為 ATS 現實說法
para(
    '現在大多數企業的第一關篩選是 ATS 系統，不是人眼——履歷在被人閱讀之前，'
    '就已經被系統根據關鍵詞過濾過一輪。通過 ATS 之後，HR 對每份履歷的人工初審時間極短，'
    '多數在 10–30 秒內決定是否繼續看，'
    '這意味著前三條工作成果的句子，決定了後面有沒有機會被讀到。',
    space_before=6
)
para('在這短短幾秒裡，決定「留下來細看」和「直接略過」的，往往不是你的學歷或年資，而是你的句型。')

# 修正 2：補選題邏輯
para(
    '這份精華版整理了 5 個關鍵句型中的 2 個。'
    '選用句型一和句型三，是因為這兩個能處理最常見的兩種問題：'
    '沒有數字的描述型句子（句型一），以及缺乏背景脈絡的成果句子（句型三）。'
    '每一個都有一個具體的改寫公式，照著填空，你的履歷就會開始說話。'
)
divider()

# ══════════════════════════════════════════════════
#  句型一：成果量化句型
# ══════════════════════════════════════════════════
heading1('句型一：成果量化句型')

heading2('為什麼「負責」這兩個字是履歷殺手？')
para('打開一百份履歷，你會發現有九十份都長這樣：', space_before=4)
bullet('負責社群媒體管理')
bullet('負責客戶服務')
bullet('協助業務開發')
para('這些句子的問題不是「假的」，而是「沒有告訴 HR 任何有用的資訊」。', space_before=6)
para(
    '「負責社群媒體管理」這句話，一個剛入職三個月的新人可以寫，一個做了五年的資深主管也可以寫。'
    'HR 看到這句話，完全無法分辨你和其他人的差異在哪裡。'
)

p_key = doc.add_paragraph()
p_key.paragraph_format.space_before = Pt(6)
p_key.paragraph_format.space_after  = Pt(4)
run = p_key.add_run('成果量化句型的核心邏輯是：把你「做了什麼」變成「創造了什麼」。')
set_font(run, 11, bold=True, color=DARK)

para('這不是造假，而是把你已經做到的事情，用一個讓人「看得見影響力」的方式說出來。')

# 修正 5：補 ATS 關聯性
para(
    '這個句型之所以在 ATS 篩選盛行的環境裡更關鍵，是因為行動動詞和具體工具名稱，'
    '同時也是 ATS 比對 JD 的高頻關鍵詞——寫得具體，人和系統都更容易看懂你。',
    color=GRAY, size=10
)

heading2('句型公式')
formula_box('【行動動詞】＋【具體方法/工具】＋【量化結果（數字/比較/百分比）】')

heading3('重點說明：')
bullet('行動動詞：用主動動詞開頭，例如：主導、設計、建立、優化、推動、整合')
bullet('具體方法：你用了什麼方式達成？讓 HR 知道你的手段')
bullet('量化結果：能加數字就加數字。沒有數字，也可以用「規模」「頻率」「範圍」描述')

heading2('改寫對照')
# 修正 6：補護理師非辦公室背景範例（第 4 行）
add_comparison_table([
    (1, '負責社群媒體管理，提升品牌知名度',
        '主導 Instagram 帳號重整，制定每週 5 篇固定發文節奏，3 個月使粉絲從 2,100 成長至 8,400（+300%），互動率從 0.8% 提升至 4.2%'),
    (2, '協助業務開發，業績良好',
        '獨立開發 B2B 新客戶，半年簽約 12 家，貢獻新增年化合約金額 NT$180 萬，達成率 120%'),
    (3, '負責人事行政工作',
        '建立新進員工到職流程，將入職培訓時間從 2 週縮短至 5 天，年省人力成本估計 NT$24 萬 *'),
    (4, '負責病患照護與衛教（護理師）',
        '設計個別化衛教追蹤表單，每日輔導 20+ 名糖尿病患者，3 個月後自主血糖管理達標率從 54% 提升至 79%'),
])

# 修正 4：人事行政 NT$24 萬估算邏輯
add_footnote(
    '* 估算基準：縮短 9 個培訓天 × 每年約 12 名新進人員 × 平均日薪成本 NT$2,000 ≈ NT$216,000（約 NT$24 萬）'
)

heading2('練習格式')
add_practice_box([
    '我做的事：_______________________________________________',
    '',
    '我用的方法：_____________________________________________',
    '',
    '結果是（數字/規模/頻率）：______________________________',
    '',
    '改寫後的句子：',
    '【行動動詞】＋________________＋________________（數字）',
])

# 修正 3：沒有數字怎麼辦 小框
add_no_number_box()

divider()

# ══════════════════════════════════════════════════
#  句型三：PAR 結構句型
# ══════════════════════════════════════════════════
heading1('句型三：PAR 結構句型（問題→行動→結果）')

heading2('為什麼說「我優化了流程」沒有說服力？')
para(
    '很多人寫履歷喜歡寫「優化」「改善」「提升效率」這類詞彙，'
    '但這些詞彙有一個根本問題：沒有背景，沒有說服力。',
    space_before=4
)
para(
    '「我優化了報表流程」這句話，HR 的第一個問題是：'
    '什麼樣的報表流程？為什麼需要優化？你怎麼優化？優化了多少？'
)

p_key2 = doc.add_paragraph()
p_key2.paragraph_format.space_before = Pt(6)
p_key2.paragraph_format.space_after  = Pt(6)
run = p_key2.add_run(
    'PAR 結構句型的核心邏輯是：讓 HR 跟著你經歷一個完整的故事——'
    '看到問題，看到你的決策，看到你創造的結果。'
)
set_font(run, 11, bold=True, color=DARK)

para('這個句型特別能展現三件事：主動性（不是被動等指示）、解決問題的能力、對成果負責的態度。')

heading2('句型公式')
formula_box('發現【問題/機會/痛點】，主導/設計/提出【具體行動】，達成【可量化結果】')

heading3('重點說明：')
bullet('問題：描述當時的現狀或痛點，讓 HR 理解背景')
bullet('行動：你做了什麼？要有主詞（我主導、我提出、我設計）')
bullet('結果：最終發生了什麼改變？用數字或具體描述')

heading2('改寫對照')
add_comparison_table([
    (1, '優化了部門報表流程',
        '發現月報需耗費 3 名同仁各 2 天手動彙整，主導建立 Excel 自動化彙報模板，將製作時間從 6 人天縮短至 4 小時，年省人力成本估算約 NT$14 萬 †'),
    (2, '改善客訴處理流程',
        '發現客訴平均回應時間長達 4.2 天，導致客戶流失率偏高；重新設計分級回應 SOP，並培訓 8 名客服同仁，3 個月後平均回應時間降至 1.5 天，客戶留存率從 72% 提升至 90%'),
    (3, '推動跨部門協作',
        '識別業務與技術部門溝通落差導致交期延誤，發起雙週跨部門同步會議，制定標準需求規格書格式，專案交期達成率從 61% 提升至 89%'),
])
add_footnote(
    '† 估算基準：6人天/月 × 12個月 = 72人天/年 × 平均日薪成本 NT$2,000 = NT$144,000，約 NT$14 萬'
)

heading2('練習格式')
add_practice_box([
    '當時遇到的問題/發現的機會是：',
    '_____________________________________________________',
    '',
    '我做了什麼（具體行動）：',
    '_____________________________________________________',
    '',
    '最終結果是：',
    '_____________________________________________________',
    '',
    '合成後的 PAR 句子：',
    '發現【問題】→ 主導【行動】→ 達成【結果（數字）】',
])

# PAR 也補沒有數字怎麼辦（同樣適用）
add_no_number_box()

divider()

# ══════════════════════════════════════════════════
#  升級 CTA（修正 7）
# ══════════════════════════════════════════════════
heading1('想要全部 5 個句型？')
para('這份精華版包含最常使用的 2 個句型。完整版還有：', space_before=6)
bullet('句型二：跨界翻譯句型 — 跨產業求職必用，用目標產業的語言重新包裝你的能力')
bullet('句型四：職涯轉折正面化句型 — 有空窗期或短暫工作記錄？用這個句型主動掌握敘事權')
bullet('句型五：JD 關鍵詞鏡像句型 — 通過 ATS 自動篩選系統的關鍵技巧')

p_cta = doc.add_paragraph()
p_cta.paragraph_format.space_before = Pt(10)
p_cta.paragraph_format.space_after  = Pt(4)
pPr = p_cta._p.get_or_add_pPr()
shd_cta = OxmlElement('w:shd')
shd_cta.set(qn('w:val'), 'clear')
shd_cta.set(qn('w:color'), 'auto')
shd_cta.set(qn('w:fill'), 'EBF4FB')
pPr.append(shd_cta)
p_cta.paragraph_format.left_indent = Cm(0.5)
run = p_cta.add_run('完整版（NT$199）加 LINE 洽詢：https://lin.ee/IOX6V66')
set_font(run, 11, bold=True, color=ACCENT)

# 修正 7：補 NT$199 與主指南定位說明
para(
    '完整句型版（NT$199）是快速操作工具，解決「怎麼說」的問題；'
    '如果你還在評估方向定位或有轉職需求，歡迎一對一諮詢——'
    '那個層次解決的是「說什麼、說給誰聽」的策略問題，兩者各自獨立。',
    size=9.5, color=GRAY, space_before=6
)

divider()

# ══════════════════════════════════════════════════
#  結語
# ══════════════════════════════════════════════════
heading1('結語')
para('這 2 個句型幫你解決的是「表達層」的問題：如何讓已有的成果被看見。', space_before=6)
para(
    '如果改完句型，投了幾封仍然沒有回音，問題通常出在「策略層」——'
    '也就是說，不只是「怎麼說」，還有「說什麼」和「說給誰聽」。'
)
para('這個部分每個人的狀況都不同，文字很難完全解決。如果你想針對自己的情況做更深入的評估，歡迎聯絡我：')
# Fix E：AI 診斷 CTA 最顯眼（轉換障礙最低，獨立框出）
shaded_para('▶ 立即試用（免費）：AI 履歷快速診斷', 'E8F4FD',
            size=12, bold=True, color=ACCENT, space_before=10, space_after=2)
shaded_para('https://resume-diagnosis.vercel.app/', 'E8F4FD',
            size=10.5, color=ACCENT, space_before=0, space_after=8)
para('📌 加 LINE 洽談：https://lin.ee/IOX6V66', space_before=0, size=10, color=GRAY)
para('📌 官網預約：https://www.careerssl.com/', space_before=0, size=10, color=GRAY)

divider()
para('職涯停看聽 | 蒲朝棟 Tim', size=10, color=GRAY,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10)
para('CDA 認證職涯顧問・104 職涯引導師', size=10, color=GRAY,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0)
para('服務 300+ 位求職者，專注 3-10 年中階職場人才', size=10, color=GRAY,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0)

# ── 儲存（雙路徑：主檔 + v3 命名版）──────────────
import os
base = r'C:\Users\USER\Desktop\職涯停看聽_網站\電子書\5句型電子書'
out_main = os.path.join(base, '讓履歷被看見的關鍵句型_免費精華版.docx')
out_v3   = os.path.join(base, '讓履歷被看見的關鍵句型_免費精華版_v3_第二次評估修正後.docx')
doc.save(out_main)
doc.save(out_v3)
print(f'✅ 主檔：{out_main}')
print(f'✅ v3 命名：{out_v3}')
