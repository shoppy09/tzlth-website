"""
產生「免費精華版 v1 原版」— 修改前版本，用於前後比對
差異：7秒說法 / 無選題邏輯 / 無沒有數字小框 / 3個範例（無護理師）/ 無ATS說明 / 無估算邏輯 / 無定位說明
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
section = doc.sections[0]
section.page_width=Cm(21); section.page_height=Cm(29.7)
section.left_margin=Cm(2.5); section.right_margin=Cm(2.5)
section.top_margin=Cm(2.5); section.bottom_margin=Cm(2.5)

DARK=RGBColor(0x1A,0x1A,0x2E); ACCENT=RGBColor(0x0F,0x72,0xB1)
GRAY=RGBColor(0x60,0x60,0x60); WHITE=RGBColor(0xFF,0xFF,0xFF)
RED=RGBColor(0xC0,0x39,0x2B); GREEN=RGBColor(0x1A,0x7A,0x4A)

def set_font(run, size, bold=False, color=None, italic=False):
    run.font.name = "微軟正黑體"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color: run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rF = OxmlElement('w:rFonts')
    rF.set(qn('w:eastAsia'), '微軟正黑體')
    rPr.insert(0, rF)

def para(text, size=11, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
         space_before=0, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(size * 1.6)
    run = p.add_run(text)
    set_font(run, size, bold, color or DARK, italic)
    return p

def heading1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), '0F72B1')
    pPr.append(shd)
    run = p.add_run(f'  {text}')
    set_font(run, 16, bold=True, color=WHITE)
    p.paragraph_format.line_spacing = Pt(28)

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, 13, bold=True, color=ACCENT)

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, 11, bold=True, color=DARK)

def formula_box(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.right_indent = Cm(0.5)
    pPr = p._p.get_or_add_pPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'EBF4FB')
    pPr.append(shd)
    run = p.add_run(text)
    set_font(run, 11, bold=True, color=ACCENT)
    p.paragraph_format.line_spacing = Pt(20)

def bullet(text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size, color=DARK)

def add_comparison_table(rows_data):
    table = doc.add_table(rows=1+len(rows_data), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for cell, txt, fill in zip(hdr, ['#', '改前 ❌', '改後 ✅'], ['2C3E50', 'C0392B', '1A7A4A']):
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(txt)
        set_font(run, 10, bold=True, color=WHITE)
        tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
        tcPr.append(shd)
    widths = [Cm(0.8), Cm(5.5), Cm(9.5)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr(); tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(widths[i].pt * 20))); tcW.set(qn('w:type'), 'dxa'); tcPr.append(tcW)
    for i, (num, before, after) in enumerate(rows_data):
        row = table.rows[i+1]
        for cell, txt, bold, color in zip(row.cells, [str(num), before, after],
                                          [True, False, False], [DARK, RED, GREEN]):
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(txt)
            set_font(run, 9.5, bold=bold, color=color)
            cell.paragraphs[0].paragraph_format.line_spacing = Pt(15)
    doc.add_paragraph()

def add_practice_box(lines):
    para('', size=4, space_before=0, space_after=2)
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F8F9FA')
    pPr.append(shd)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    run = p.add_run('✏️ 練習格式')
    set_font(run, 10, bold=True, color=ACCENT)
    for line in lines:
        p2 = doc.add_paragraph()
        pPr2 = p2._p.get_or_add_pPr(); shd2 = OxmlElement('w:shd')
        shd2.set(qn('w:val'), 'clear'); shd2.set(qn('w:color'), 'auto'); shd2.set(qn('w:fill'), 'F8F9FA')
        pPr2.append(shd2)
        p2.paragraph_format.left_indent = Cm(0.5)
        p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(3)
        run2 = p2.add_run(line)
        set_font(run2, 10, color=DARK); p2.paragraph_format.line_spacing = Pt(18)

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    pBdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom); p._p.get_or_add_pPr().append(pBdr)

# ═══ 封面 ═══
p_c = doc.add_paragraph()
p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_c.paragraph_format.space_before = Pt(60); p_c.paragraph_format.space_after = Pt(8)
run = p_c.add_run('讓履歷被看見的關鍵句型')
set_font(run, 28, bold=True, color=ACCENT)
para('免費精華版', size=14, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, space_before=0, space_after=4)
para('2 個最高頻使用的改寫公式，附改寫前後對照', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, space_before=0, space_after=40)
para('蒲朝棟 Tim', size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=DARK, space_before=0, space_after=4)
para('CDA 認證職涯顧問・104 職涯引導師', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, space_before=0, space_after=60)
doc.add_page_break()

# ═══ 給你的一段話（v1：含7秒，無ATS，無選題邏輯）═══
heading1('給你的一段話')
para('HR 每天要看幾十份、甚至幾百份履歷，平均花在每份上的時間不超過 7 秒。', space_before=6)
para('在這 7 秒裡，決定「留下來細看」和「直接略過」的，往往不是你的學歷或年資，而是你的句型。')
para('這份精華版整理了 5 個關鍵句型中最高頻使用的 2 個。每一個都有一個具體的改寫公式，你不需要是文字高手，照著公式填空，你的履歷就會開始說話。')
divider()

# ═══ 句型一（v1：3個範例，無ATS說明，無護理師，無小框）═══
heading1('句型一：成果量化句型')
heading2('為什麼「負責」這兩個字是履歷殺手？')
para('打開一百份履歷，你會發現有九十份都長這樣：', space_before=4)
bullet('負責社群媒體管理'); bullet('負責客戶服務'); bullet('協助業務開發')
para('這些句子的問題不是「假的」，而是「沒有告訴 HR 任何有用的資訊」。', space_before=6)
para('「負責社群媒體管理」這句話，一個剛入職三個月的新人可以寫，一個做了五年的資深主管也可以寫。HR 看到這句話，完全無法分辨你和其他人的差異在哪裡。')
p_k = doc.add_paragraph()
p_k.paragraph_format.space_before = Pt(6); p_k.paragraph_format.space_after = Pt(6)
run = p_k.add_run('成果量化句型的核心邏輯是：把你「做了什麼」變成「創造了什麼」。')
set_font(run, 11, bold=True, color=DARK)
para('這不是造假，而是把你已經做到的事情，用一個讓人「看得見影響力」的方式說出來。')
heading2('句型公式')
formula_box('【行動動詞】＋【具體方法/工具】＋【量化結果（數字/比較/百分比）】')
heading3('重點說明：')
bullet('行動動詞：用主動動詞開頭，例如：主導、設計、建立、優化、推動、整合')
bullet('具體方法：你用了什麼方式達成？讓 HR 知道你的手段')
bullet('量化結果：能加數字就加數字。沒有數字，也可以用「規模」「頻率」「範圍」描述')
heading2('改寫對照')
add_comparison_table([
    (1, '負責社群媒體管理，提升品牌知名度',
     '主導 Instagram 帳號重整，制定每週 5 篇固定發文節奏，3 個月使粉絲從 2,100 成長至 8,400（+300%），互動率從 0.8% 提升至 4.2%'),
    (2, '協助業務開發，業績良好',
     '獨立開發 B2B 新客戶，半年簽約 12 家，貢獻新增年營收 NT$180 萬，達成率 120%'),
    (3, '負責人事行政工作',
     '建立新進員工到職流程，將入職培訓時間從 2 週縮短至 5 天，年省人力成本估計 NT$24 萬'),
])
heading2('練習格式')
add_practice_box([
    '我做的事：_______________________________________________', '',
    '我用的方法：_____________________________________________', '',
    '結果是（數字/規模/頻率）：______________________________', '',
    '改寫後的句子：', '【行動動詞】＋________________＋________________（數字）',
])
divider()

# ═══ 句型三（v1 原版，無小框）═══
heading1('句型三：PAR 結構句型（問題→行動→結果）')
heading2('為什麼說「我優化了流程」沒有說服力？')
para('很多人寫履歷喜歡寫「優化」「改善」「提升效率」這類詞彙，但這些詞彙有一個根本問題：沒有背景，沒有說服力。', space_before=4)
para('「我優化了報表流程」這句話，HR 的第一個問題是：什麼樣的報表流程？為什麼需要優化？你怎麼優化？優化了多少？')
p_k2 = doc.add_paragraph()
p_k2.paragraph_format.space_before = Pt(6); p_k2.paragraph_format.space_after = Pt(6)
run = p_k2.add_run('PAR 結構句型的核心邏輯是：讓 HR 跟著你經歷一個完整的故事——看到問題，看到你的決策，看到你創造的結果。')
set_font(run, 11, bold=True, color=DARK)
para('這個句型特別能展現三件事：主動性（不是被動等指示）、解決問題的能力、對成果負責的態度。')
heading2('句型公式')
formula_box('發現【問題/機會/痛點】，主導/設計/提出【具體行動】，達成【可量化結果】')
heading3('重點說明：')
bullet('問題：描述當時的現狀或痛點，讓 HR 理解背景')
bullet('行動：你做了什麼？要有主詞（我主導、我提出、我設計）')
bullet('結果：最終發生了什麼改變？用數字或具體描述')
heading2('改寫對照')
add_comparison_table([
    (1, '優化了部門報表流程',
     '發現月報需耗費 3 名同仁各 2 天手動彙整，主導建立 Excel 自動化彙報模板，將製作時間從 6 人天縮短至 4 小時，年省人力成本估算約 NT$30 萬'),
    (2, '改善客訴處理流程',
     '發現客訴平均回應時間長達 4.2 天，導致客戶流失率偏高；重新設計分級回應 SOP，並培訓 8 名客服同仁，3 個月後平均回應時間降至 1.5 天，客戶留存率提升 18%'),
    (3, '推動跨部門協作',
     '識別業務與技術部門溝通落差導致交期延誤，發起雙週跨部門同步會議，制定標準需求規格書格式，專案交期達成率從 61% 提升至 89%'),
])
doc.add_paragraph()
heading2('練習格式')
add_practice_box([
    '當時遇到的問題/發現的機會是：', '_____________________________________________________', '',
    '我做了什麼（具體行動）：', '_____________________________________________________', '',
    '最終結果是：', '_____________________________________________________', '',
    '合成後的 PAR 句子：', '發現【問題】→ 主導【行動】→ 達成【結果（數字）】',
])
divider()

# ═══ CTA（v1：無定位說明）═══
heading1('想要全部 5 個句型？')
para('這份精華版包含最常使用的 2 個句型。完整版還有：', space_before=6)
bullet('句型二：跨界翻譯句型 — 跨產業求職必用，用目標產業的語言重新包裝你的能力')
bullet('句型四：職涯轉折正面化句型 — 有空窗期或短暫工作記錄？用這個句型主動掌握敘事權')
bullet('句型五：JD 關鍵詞鏡像句型 — 通過 ATS 自動篩選系統的關鍵技巧')
p_cta = doc.add_paragraph()
p_cta.paragraph_format.space_before = Pt(10); p_cta.paragraph_format.space_after = Pt(4)
pPr = p_cta._p.get_or_add_pPr(); shd_cta = OxmlElement('w:shd')
shd_cta.set(qn('w:val'), 'clear'); shd_cta.set(qn('w:color'), 'auto'); shd_cta.set(qn('w:fill'), 'EBF4FB')
pPr.append(shd_cta); p_cta.paragraph_format.left_indent = Cm(0.5)
run = p_cta.add_run('完整版（NT$199）加 LINE 洽詢：https://lin.ee/IOX6V66')
set_font(run, 11, bold=True, color=ACCENT)
divider()

# ═══ 結語 ═══
heading1('結語')
para('這 2 個句型幫你解決的是「表達層」的問題：如何讓已有的成果被看見。', space_before=6)
para('如果改完句型，投了幾封仍然沒有回音，問題通常出在「策略層」——也就是說，不只是「怎麼說」，還有「說什麼」和「說給誰聽」。')
para('這個部分每個人的狀況都不同，文字很難完全解決。如果你想針對自己的情況做更深入的評估，歡迎聯絡我：')
para('📌 加 LINE 洽談：https://lin.ee/IOX6V66', space_before=6)
para('📌 官網預約：https://www.careerssl.com/', space_before=0)
para('📌 AI 履歷快速診斷（免費）：https://resume-diagnosis.vercel.app/', space_before=0)
divider()
para('職涯停看聽 | 蒲朝棟 Tim', size=10, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10)
para('CDA 認證職涯顧問・104 職涯引導師', size=10, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0)
para('服務 300+ 位求職者，專注 3-10 年中階職場人才', size=10, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0)

out = r'C:\Users\USER\Desktop\職涯停看聽_網站\電子書\5句型電子書\讓履歷被看見的關鍵句型_免費精華版_v1_原版.docx'
doc.save(out)
print(f'v1 saved: {out}')
