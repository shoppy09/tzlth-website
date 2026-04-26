"""
產生「讓履歷被看見的5個關鍵句型_完整版.docx」
v5 全面改版（2026-04-26）- 深度審查報告 11問題 + 3破局點
  v1-v4 修正（見前版本記錄）
  v5 新增修正（深度審查報告）：
    Fix 16：引言 開場三連句→去平行句式版（保留衝擊感）
    Fix 17：引言 新增表達層vs策略層框架說明（破局點二）
    Fix 18：使用說明 比較表改版（適合你是/解決的問題格式）
    Fix 19：使用說明後 新增決策樹（破局點一）
    Fix 20：句型一 新增「如何找到自己的數字」四管道
    Fix 21：句型一 新增三步驟量化公式（破局點三）
    Fix 22：句型二 適用情境框加句型二vs五差異說明
    Fix 23：句型三開頭 新增句型一vs三差異說明
    Fix 24：句型三 第二次「沒有數字怎麼辦」→一行引導
    Fix 25：句型四例4改後 替換為語氣更真實的版本
    Fix 26：句型五 適用情境框加句型二vs五差異說明
    Fix 27：句型五 練習格式→三步驟+自我檢核表格式
    Fix 28：結語 後半段AI腔→真實案例語氣版本
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── 頁面設定 ──────────────────────────────────────
section = doc.sections[0]
section.page_width    = Cm(21)
section.page_height   = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── 顏色常數 ──────────────────────────────────────
DARK   = RGBColor(0x1A, 0x1A, 0x2E)
ACCENT = RGBColor(0x0F, 0x72, 0xB1)
GRAY   = RGBColor(0x60, 0x60, 0x60)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
RED    = RGBColor(0xC0, 0x39, 0x2B)
GREEN  = RGBColor(0x1A, 0x7A, 0x4A)
ORANGE = RGBColor(0xE6, 0x7E, 0x22)

# ── 工具函式 ──────────────────────────────────────
def set_font(run, size, bold=False, color=None, italic=False):
    run.font.name = '微軟正黑體'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    rPr.insert(0, rFonts)

def para(text, size=11, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
         space_before=0, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(size * 1.6)
    run = p.add_run(text)
    set_font(run, size, bold, color or DARK, italic)
    return p

def heading1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '0F72B1')
    pPr.append(shd)
    run = p.add_run(f'  {text}')
    set_font(run, 16, bold=True, color=WHITE)
    p.paragraph_format.line_spacing = Pt(28)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, 13, bold=True, color=ACCENT)
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, 11, bold=True, color=DARK)
    return p

def formula_box(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EBF4FB')
    pPr.append(shd)
    run = p.add_run(text)
    set_font(run, 11, bold=True, color=ACCENT)
    p.paragraph_format.line_spacing = Pt(20)
    return p

def bullet(text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size, color=DARK)
    return p

def shaded_para(text, fill_hex, size=10, bold=False, color=None,
                left_indent=0.5, space_before=0, space_after=3):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)
    p.paragraph_format.left_indent  = Cm(left_indent)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(size * 1.65)
    run = p.add_run(text)
    set_font(run, size, bold=bold, color=color or DARK)
    return p

def add_tip_box(title, lines, fill='FFF3CD', title_color=None):
    tc = title_color or ORANGE
    shaded_para(f'💡 {title}', fill, size=10, bold=True, color=tc,
                space_before=8, space_after=2)
    for line in lines:
        shaded_para(line, fill, size=9.5, color=DARK, space_before=0, space_after=3)

def add_footnote(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.3)
    run = p.add_run(text)
    set_font(run, 8.5, color=GRAY, italic=True)

def add_comparison_table(rows_data, col_widths=None):
    table = doc.add_table(rows=1 + len(rows_data), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for cell, txt, fill in zip(hdr, ['#', '改前 ❌', '改後 ✅'],
                                ['2C3E50', 'C0392B', '1A7A4A']):
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(txt)
        set_font(run, 10, bold=True, color=WHITE)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        tcPr.append(shd)
    widths = col_widths or [Cm(0.8), Cm(5.5), Cm(9.5)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(widths[i].pt * 20)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
    for i, (num, before, after) in enumerate(rows_data):
        row = table.rows[i + 1]
        for cell, txt, bold, color in zip(
                row.cells, [str(num), before, after],
                [True, False, False], [DARK, RED, GREEN]):
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(txt)
            set_font(run, 9.5, bold=bold, color=color)
            cell.paragraphs[0].paragraph_format.line_spacing = Pt(15)

def add_comparison_table_4col(rows_data):
    """4欄表格（句型二：原始背景/目標職位/改前/改後）"""
    table = doc.add_table(rows=1 + len(rows_data), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    headers = ['#', '原始背景', '改前 ❌', '改後 ✅']
    fills   = ['2C3E50', '34495E', 'C0392B', '1A7A4A']
    for cell, txt, fill in zip(hdr, headers, fills):
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(txt)
        set_font(run, 9.5, bold=True, color=WHITE)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        tcPr.append(shd)
    widths = [Cm(0.6), Cm(3.0), Cm(4.5), Cm(7.7)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(widths[i].pt * 20)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
    for i, (num, bg, before, after) in enumerate(rows_data):
        row = table.rows[i + 1]
        data = [str(num), bg, before, after]
        colors = [DARK, DARK, RED, GREEN]
        bolds  = [True, False, False, False]
        for cell, txt, bold, color in zip(row.cells, data, bolds, colors):
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(txt)
            set_font(run, 9, bold=bold, color=color)
            cell.paragraphs[0].paragraph_format.line_spacing = Pt(14)

def add_comparison_table_jd(rows_data):
    """JD 關鍵詞表格（4欄：#/JD關鍵詞/改前/改後）"""
    table = doc.add_table(rows=1 + len(rows_data), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    headers = ['#', 'JD 關鍵詞', '改前 ❌', '改後 ✅（關鍵詞已嵌入）']
    fills   = ['2C3E50', '5D6D7E', 'C0392B', '1A7A4A']
    for cell, txt, fill in zip(hdr, headers, fills):
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(txt)
        set_font(run, 9.5, bold=True, color=WHITE)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        tcPr.append(shd)
    widths = [Cm(0.6), Cm(3.5), Cm(4.0), Cm(7.7)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(widths[i].pt * 20)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
    for i, (num, kw, before, after) in enumerate(rows_data):
        row = table.rows[i + 1]
        data = [str(num), kw, before, after]
        colors = [DARK, ACCENT, RED, GREEN]
        bolds  = [True, True, False, False]
        for cell, txt, bold, color in zip(row.cells, data, bolds, colors):
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(txt)
            set_font(run, 9, bold=bold, color=color)
            cell.paragraphs[0].paragraph_format.line_spacing = Pt(14)

def add_comparison_table_situation(rows_data):
    """情境4欄表格（句型四：#/情境/改前/改後）"""
    table = doc.add_table(rows=1 + len(rows_data), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    headers = ['#', '情境', '改前 ❌', '改後 ✅']
    fills   = ['2C3E50', '5D6D7E', 'C0392B', '1A7A4A']
    for cell, txt, fill in zip(hdr, headers, fills):
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(txt)
        set_font(run, 9.5, bold=True, color=WHITE)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        tcPr.append(shd)
    widths = [Cm(0.6), Cm(3.0), Cm(4.2), Cm(8.0)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(widths[i].pt * 20)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
    for i, (num, sit, before, after) in enumerate(rows_data):
        row = table.rows[i + 1]
        data = [str(num), sit, before, after]
        colors = [DARK, DARK, RED, GREEN]
        bolds  = [True, False, False, False]
        for cell, txt, bold, color in zip(row.cells, data, bolds, colors):
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(txt)
            set_font(run, 9, bold=bold, color=color)
            cell.paragraphs[0].paragraph_format.line_spacing = Pt(14)

def add_practice_box(lines):
    para('', size=4, space_before=0, space_after=2)
    shaded_para('✏️ 練習格式', 'F8F9FA', size=10, bold=True, color=ACCENT,
                space_before=6, space_after=2)
    for line in lines:
        shaded_para(line, 'F8F9FA', size=10, color=DARK,
                    space_before=0, space_after=3)

def add_no_number_box():
    add_tip_box(
        '沒有數字怎麼辦？',
        [
            '如果無法取得精確數字，可以用以下方式替代：',
            '・規模（例：服務 50+ 位客戶、管理 3 人小組）',
            '・頻率（例：每週主持 3 場跨部門會議）',
            '・相對比較（例：從每週 2 次縮短至每週 1 次以下）',
            '・定性影響（例：首次建立 XX 機制，此前無標準流程）',
            '',
            '有方向的估算，比空白好。只要推算邏輯合理，估算值也可以寫進履歷。',
        ],
        fill='FFF8E1',
        title_color=ORANGE,
    )

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    p._p.get_or_add_pPr().append(pBdr)

def add_hr_visual():
    shaded_para(
        '以下範例數字為真實輔導案例的示意數字。你的數字不需要和範例一樣大，只需要是真實的、有邏輯支撐的。',
        'FEF9E7', size=9.5, color=DARK, space_before=8, space_after=2
    )
    shaded_para(
        '改前 HR 停留機率：⬛⬜⬜⬜⬜（約 1/5）｜改後：⬛⬛⬛⬛⬛（約 5/5）',
        'FEF9E7', size=10, bold=True, color=DARK, space_before=0, space_after=4
    )

# ══════════════════════════════════════════════════
#  封面頁
# ══════════════════════════════════════════════════
p_cover = doc.add_paragraph()
p_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_cover.paragraph_format.space_before = Pt(60)
p_cover.paragraph_format.space_after  = Pt(8)
run = p_cover.add_run('你的履歷沒有說謊，但它說錯話了')
set_font(run, 26, bold=True, color=ACCENT)

para('完整版', size=14, align=WD_ALIGN_PARAGRAPH.CENTER,
     color=GRAY, space_before=0, space_after=4)
para('從「負責執行」到「創造成果」的 5 個改寫公式',
     size=11, align=WD_ALIGN_PARAGRAPH.CENTER,
     color=GRAY, space_before=0, space_after=40)
para('蒲朝棟 Tim', size=12, bold=True,
     align=WD_ALIGN_PARAGRAPH.CENTER, color=DARK, space_before=0, space_after=4)
para('CDA 認證職涯顧問・104 職涯引導師',
     size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, space_before=0, space_after=60)

doc.add_page_break()

# ══════════════════════════════════════════════════
#  給你的一段話
# ══════════════════════════════════════════════════
heading1('給你的一段話')

para('每次幫求職者看履歷，我最常說的一句話是：', space_before=6)

p_quote = doc.add_paragraph()
p_quote.paragraph_format.space_before = Pt(4)
p_quote.paragraph_format.space_after  = Pt(4)
p_quote.paragraph_format.left_indent  = Cm(1.0)
run = p_quote.add_run('「你做的事情很好，但你沒有讓它被看見。」')
set_font(run, 12, bold=True, color=ACCENT)
p_quote.paragraph_format.line_spacing = Pt(22)

para('不是能力不夠，是語言出了問題。')
para(
    'HR 每天要看幾十份、甚至幾百份履歷，人工初審時間極短，多數在 10–30 秒內決定是否繼續看。'
    '在這短短幾秒裡，決定「留下來細看」和「直接略過」的，往往不是你的學歷或年資，而是你的句型。'
)
para(
    '大多數求職者不知道的是：HR 初審時不是在「評分」你，他們是在「尋找信號」——'
    '有沒有成果？成果是否具體？能不能快速讀懂？句型寫對了，你就是那個信號；'
    '寫錯了，再好的能力也會被跳過。'
)
para(
    '說明一點：這本手冊處理的是表達層的問題——你做了這些事，但還沒有讓它被看見。'
    '如果你改完句型後發現投件仍然沒有回音，問題可能在策略層（投的方向對不對、'
    '你的定位清不清楚）——那超出了句型能處理的範圍。結語裡我會說明兩者的差異，'
    '以及下一步可以怎麼做。',
    color=GRAY, size=10
)
para(
    '這本手冊整理了 5 個我在諮詢中反覆使用的關鍵句型。'
    '每一個都有一個具體的改寫公式，你不需要是文字高手，'
    '照著公式填空，你的履歷就會開始說話。'
)
divider()

# ══════════════════════════════════════════════════
#  使用說明
# ══════════════════════════════════════════════════
heading1('使用說明')
heading2('這本手冊適合你，如果你：')
bullet('寫了很多年的履歷，但從來不知道哪裡寫錯了')
bullet('覺得自己「做了很多」，但寫出來就是不吸引人')
bullet('正在轉職、有空窗期、或投了幾十封都沒有面試邀請')

heading2('如何使用：')
bullet('每個句型先讀「為什麼這樣寫」，理解背後邏輯')
bullet('看改寫對照，感受「前後差距」')
bullet('用書末的練習格式，套用在你自己的經歷上')

# Fix 18：使用說明對比表改版（你是誰/解決的問題格式）
heading2('完整版 vs 免費精華版')
table_cmp = doc.add_table(rows=5, cols=3)
table_cmp.alignment = WD_TABLE_ALIGNMENT.CENTER
table_cmp.style = 'Table Grid'
cmp_hdr = table_cmp.rows[0].cells
for cell, txt, fill in zip(cmp_hdr, ['', '免費精華版', '完整版（本冊）'],
                            ['2C3E50', '5D6D7E', '0F72B1']):
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(txt)
    set_font(run, 9.5, bold=True, color=WHITE)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)
cmp_rows = [
    ('適合你是', '有工作成果但不知如何量化，或有空窗期、短暫工作需要說明的求職者', '含跨界、空窗、頻繁換工作、ATS 卡關等複雜背景的求職者'),
    ('解決的問題', '如何把成果用數字說清楚，以及如何主動說明職涯轉折，讓履歷完整可信', '如何讓 HR 看懂你的背景，並通過多種篩選關卡'),
    ('句型數量', '2 個（句型一、句型四）', '5 個，完整涵蓋'),
    ('改寫範例數', '各 3–4 例', '各 3–4 例，部分情境額外擴充'),
]
for i, (label, free, full) in enumerate(cmp_rows):
    row = table_cmp.rows[i + 1]
    for cell, txt, bold, color in zip(row.cells,
            [label, free, full], [True, False, False],
            [DARK, GRAY, GREEN]):
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(txt)
        set_font(run, 9.5, bold=bold, color=color)
        cell.paragraphs[0].paragraph_format.line_spacing = Pt(14)

# Fix 19：決策樹
para('', size=4, space_before=0, space_after=2)
heading3('不確定從哪個句型開始？')
add_tip_box(
    '你寫的這條描述，現在長什麼樣子？',
    [
        '📌 寫的是「負責 XX 工作」，沒有成果數字 → 先看句型一（成果量化）',
        '📌 我的背景跟目標職位差很多，對方可能看不懂 → 先看句型二（跨界翻譯）',
        '📌 我做了一件有意義的事，但一句話說不清楚 → 先看句型三（PAR 結構）',
        '📌 有空窗期、短暫工作、頻繁換工作，不知道怎麼解釋 → 先看句型四（職涯轉折正面化）',
        '📌 投了很多封都沒有回音，不知道問題在哪 → 先看句型五（JD 關鍵詞鏡像）',
    ],
    fill='EBF4FB',
    title_color=ACCENT,
)
divider()

# ══════════════════════════════════════════════════
#  句型一：成果量化句型
# ══════════════════════════════════════════════════
heading1('句型一：成果量化句型')

heading2('為什麼「負責」這兩個字是履歷殺手？')
para(
    '一個剛入職三個月的新人可以寫，一個做了五年的資深主管也可以寫——'
    'HR 看到這句話，完全無法分辨你和其他人的差異在哪裡。這就是「負責」兩個字的真實代價。',
    space_before=4
)
para('打開一百份履歷，你會發現有九十份都長這樣：')
bullet('負責社群媒體管理')
bullet('負責客戶服務')
bullet('協助業務開發')
para('這些句子的問題不是「假的」，而是「沒有告訴 HR 任何有用的資訊」。', space_before=6)

p_key = doc.add_paragraph()
p_key.paragraph_format.space_before = Pt(6)
p_key.paragraph_format.space_after  = Pt(4)
run = p_key.add_run('成果量化句型的核心邏輯是：把你「做了什麼」變成「創造了什麼」。')
set_font(run, 11, bold=True, color=DARK)

para('這不是造假，而是把你已經做到的事情，用一個讓人「看得見影響力」的方式說出來。')

para(
    '這個句型之所以在 ATS 篩選盛行的環境裡更關鍵，是因為行動動詞和具體工具名稱，'
    '同時也是 ATS 比對 JD 的高頻關鍵詞——寫得具體，人和系統都更容易看懂你。',
    color=GRAY, size=10
)

heading2('句型公式')
formula_box('【行動動詞】＋【具體方法/工具】＋【量化結果（數字/比較/百分比）】')

heading3('重點說明：')
bullet('行動動詞：用主動動詞開頭，例如：主導、設計、建立、優化、推動、整合')
bullet('具體方法：你用了什麼方式達成？讓 HR 知道你的手段')
bullet('量化結果：能加數字就加數字。沒有數字，也可以用「規模」「頻率」「範圍」描述')

add_hr_visual()
heading2('改寫對照')
add_comparison_table([
    (1, '負責社群媒體管理，提升品牌知名度',
        '主導 Instagram 帳號重整，制定每週 5 篇固定發文節奏，3 個月使粉絲從 2,100 成長至 8,400（+300%），互動率從 0.8% 提升至 4.2%'),
    (2, '協助業務開發，業績良好',
        '獨立開發 B2B 新客戶，半年簽約 12 家，貢獻新增年化合約金額 NT$180 萬，達成率 120%'),
    (3, '負責人事行政工作',
        '建立新進員工到職流程，將入職培訓時間從 2 週縮短至 5 天，年省人力成本估計 NT$12 萬 *'),
    (4, '負責病患照護與衛教（護理師）',
        '設計個別化衛教追蹤表單，每日輔導 20+ 名糖尿病患者，3 個月後自主血糖管理達標率從 54% 提升至 79%'),
])
add_footnote(
    '* 估算基準：縮短 5 個培訓天 × 每年約 12 名新進人員 × 平均日薪成本 NT$2,000 = NT$120,000（約 NT$12 萬）'
)

heading2('練習格式')
add_practice_box([
    '我做的事：_______________________________________________',
    '',
    '我用的方法：_____________________________________________',
    '',
    '結果是（數字/規模/頻率）：______________________________',
    '',
    '改寫後的句子：',
    '【行動動詞】＋________________＋________________（數字）',
])
add_no_number_box()
add_tip_box(
    '改完句型一後，想確認整份履歷還有哪些調整空間？',
    ['▶ AI 履歷快速診斷（免費）：https://resume-diagnosis.vercel.app/',
     '改完句型，送進工具看 ATS 層面還有什麼缺口。'],
    fill='E8F4FD', title_color=ACCENT,
)
divider()


# Fix 20：如何找到自己的數字
heading3('如何找到自己的數字？')
para('不是沒有數字，是還沒想到去哪裡找。以下四個管道，幾乎每份工作都適用：', size=10, space_before=2)
bullet('管道一：歷史紀錄 — 工作週報、月報、績效考核表、簡報——你曾經彙整過的數字多數在這裡', size=10)
bullet('管道二：詢問前同事或主管 — 傳訊息給前同事確認記憶中的數字是否大致正確——通常對方三句話就能幫你確認，不需要正式詢問', size=10)
bullet('管道三：行業基準對照 — 查同職位的市場數據或行業平均，以此為基準描述自己的規模或成果', size=10)
bullet('管道四：合理推算 — 如「2 週縮短至 5 天」可推算「12 名新進 × 5 天 × NT$2,000 ≈ NT$12 萬」，邏輯清楚的估算值可標注後寫進履歷', size=10)

# Fix 21：三步驟量化公式
heading3('三步驟量化公式')
para('不知道怎麼「推算」？按這三步走：', size=10, space_before=2)
bullet('Step 1：找出「變化前後」 — 這件事做之前狀況是什麼？做完之後呢？', size=10)
bullet('Step 2：選定「比較維度」 — 時間（縮短幾天）、金額（省了多少錢）、比例（提升幾個百分點）、規模（影響幾個人）', size=10)
bullet('Step 3：補上「推算鏈」 — 若無精確數字，寫出計算方式，讓 HR 能跟著你的邏輯走', size=10)
add_tip_box(
    '估算不是造假',
    ['估算是「讓成果可見」的方式。只要邏輯有依據，在履歷中標注「估算值」是完全被接受的做法。'],
    fill='FFF8E1', title_color=ORANGE,
)
# ══════════════════════════════════════════════════
#  句型二：跨界翻譯句型
# ══════════════════════════════════════════════════
heading1('句型二：跨界翻譯句型')

heading2('為什麼你的經歷「讀不懂」？')
para(
    '跨產業求職最大的困境不是「能力不夠」，而是語言落差。',
    space_before=4
)
para(
    '你在舊產業說「帶領班兵完成任務」，新產業的 HR 聽到的是「軍隊的事，跟我沒關係」。'
    '你在學校說「帶領學生達成教學目標」，企業 HR 聽到的是「老師要轉職，不確定她能不能適應企業文化」。'
)
para('問題不在你的能力，在於你用了對方看不懂的語言。')

p_key2 = doc.add_paragraph()
p_key2.paragraph_format.space_before = Pt(6)
p_key2.paragraph_format.space_after  = Pt(4)
run = p_key2.add_run('跨界翻譯句型的核心邏輯是：用目標產業的語言，重新包裝你已有的能力。')
set_font(run, 11, bold=True, color=DARK)

para(
    '這需要你先做一件事：去看你目標職位的 JD，找出它最常出現的關鍵詞，'
    '然後把你的經歷「翻譯」成那個語言。'
)

heading2('句型公式')
formula_box('【原有職能】→ 用【目標產業關鍵詞】重新框架 → 【可遷移的成果】')

heading3('重點說明：')
bullet('先列出目標 JD 的前 5 個核心能力需求')
bullet('找出你過去的工作中，哪些事情「本質上」與這 5 個需求相同')
bullet('用目標產業的語言重寫，而非原產業的術語')

# Fix 6 + Fix 22：句型二 適用情境框（加句型二vs五差異）
add_tip_box(
    '【適用情境】',
    [
        '此句型特別適合：跨產業求職、轉職初期、或背景與目標職位差異較大的求職者。',
        '若你的經歷與目標職位表面上看來「不相關」，這個句型能幫你讓 HR 看懂你的可遷移價值；',
        '若你是同產業橫向移動，也能用此方法強化語言對齊，讓關鍵能力更清晰易讀。',
        '',
        '與句型五（JD 關鍵詞鏡像）的差異：句型二是「把你的背景翻譯成對方看得懂的語言」，',
        '適用於原始背景與目標職位框架不同時；句型五是「把 JD 關鍵詞嵌入你已有的成果描述」，',
        '適用於經歷本身匹配、只是用詞沒有對齊時。若兩者都需要，建議先用句型二，再用句型五。',
    ],
    fill='EBF4FB',
    title_color=ACCENT,
)

add_hr_visual()
heading2('改寫對照')
add_comparison_table_4col([
    (1, '教師 → 企業內訓',
        '擔任國中數學教師，負責班級教學',
        '設計並執行 28 人課程方案，依學習者差異化調整教學策略，學期末評量通過率從 62% 提升至 84%（+22 個百分點）'),
    (2, '軍官 → 專案管理',
        '擔任連長，帶領官兵完成訓練任務',
        '統籌 80 人跨部門任務執行，協調資源分配與時程管控，連續 3 年訓練評鑑達甲等'),
    (3, '護理師 → 醫療業務',
        '擔任護理師，負責病患照護與衛教',
        '每日面對 20+ 名客戶（病患），執行需求評估、解決方案說明與追蹤回訪，衛教滿意度達 92%'),
])

heading2('練習格式')
add_practice_box([
    '我過去的職稱/職能：______________________________________',
    '',
    '目標職位 JD 關鍵詞（列出 3 個）：',
    '1. ___  2. ___  3. ___',
    '',
    '我做過的哪件事，本質上對應了這些關鍵詞？',
    '_____________________________________________________',
    '',
    '翻譯後的句子：',
    '【用目標關鍵詞開頭】＋【你的實際行動】＋【可量化的成果】',
])
divider()

# ══════════════════════════════════════════════════
#  句型三：PAR 結構句型
# ══════════════════════════════════════════════════
heading1('句型三：PAR 結構句型（問題→行動→結果）')

# Fix 23：句型一 vs 句型三差異說明
add_tip_box(
    '句型一 vs 句型三，有什麼不同？',
    [
        '句型一的重點是「量化」——把你的成果用數字說出來，適合當成果本身清楚、可以直接用一句話說明時。',
        '句型三（PAR 結構）的重點是「敘事」——讓 HR 跟著你看到問題、你的決策，以及你創造的結果，',
        '適合當成果需要背景才能顯現意義時。',
        '',
        '判斷方式：你能用一句量化句完整說清楚 → 用句型一。需要解釋「為什麼重要、怎麼做到的」→ 用句型三。',
    ],
    fill='F0F3F4',
    title_color=DARK,
)

heading2('為什麼說「我優化了流程」沒有說服力？')
para(
    '很多人寫履歷喜歡寫「優化」「改善」「提升效率」這類詞彙，'
    '但這些詞彙有一個根本問題：沒有背景，沒有說服力。',
    space_before=4
)
para(
    '「我優化了報表流程」這句話，HR 的第一個問題是：'
    '什麼樣的報表流程？為什麼需要優化？你怎麼優化？優化了多少？'
)

p_key3 = doc.add_paragraph()
p_key3.paragraph_format.space_before = Pt(6)
p_key3.paragraph_format.space_after  = Pt(6)
run = p_key3.add_run(
    'PAR 結構句型的核心邏輯是：讓 HR 跟著你經歷一個完整的故事——'
    '看到問題，看到你的決策，看到你創造的結果。'
)
set_font(run, 11, bold=True, color=DARK)

para('這個句型特別能展現三件事：主動性（不是被動等指示）、解決問題的能力、對成果負責的態度。')

heading2('句型公式')
formula_box('發現【問題/機會/痛點】，主導/設計/提出【具體行動】，達成【可量化結果】')

heading3('重點說明：')
bullet('問題：描述當時的現狀或痛點，讓 HR 理解背景')
bullet('行動：你做了什麼？要有主詞（我主導、我提出、我設計）')
bullet('結果：最終發生了什麼改變？用數字或具體描述')

add_hr_visual()
heading2('改寫對照')
add_comparison_table([
    (1, '優化了部門報表流程',
        '發現月報需耗費 3 名同仁各 2 天手動彙整，主導建立 Excel 自動化彙報模板，將製作時間從 6 人天縮短至 4 小時，年省人力成本估算約 NT$13 萬 †'),
    (2, '改善客訴處理流程',
        '發現客訴平均回應時間長達 4.2 天，導致客訴積壓問題明顯；重新設計分級回應 SOP，並培訓 8 名客服同仁，3 個月後平均回應時間降至 1.5 天，客訴處理滿意度從 61% 提升至 84%'),
    (3, '推動跨部門協作',
        '識別業務與技術部門溝通落差導致交期延誤，發起雙週跨部門同步會議，制定標準需求規格書格式，專案交期達成率從 58% 提升至 89%'),
    (4, '管理業務團隊，達成業績目標（客服中心）',
        '識別客服團隊缺乏統一服務標準導致 NPS（淨推薦分數）偏低，帶領 5 人團隊建立服務規範並月度校準績效，3 個月後 NPS（淨推薦分數）從 64 提升至 80，人員流動率從 42% 降至 18%'),
])
add_footnote(
    '† 估算基準：節省工時 (6 − 0.5) 人天/月 × 12 個月 = 66 人天/年 × 平均日薪成本 NT$2,000 = NT$132,000，約 NT$13 萬'
)

heading2('練習格式')
add_practice_box([
    '當時遇到的問題/發現的機會是：',
    '_____________________________________________________',
    '',
    '我做了什麼（具體行動）：',
    '_____________________________________________________',
    '',
    '最終結果是：',
    '_____________________________________________________',
    '',
    '合成後的 PAR 句子：',
    '發現【問題】→ 主導【行動】→ 達成【結果（數字）】',
])
# Fix 24：第二次「沒有數字怎麼辦」改為一行引導
add_tip_box(
    '找不到數字？',
    ['估算邏輯與四個管道，詳見句型一「沒有數字怎麼辦」與「三步驟量化公式」。'],
    fill='FFF8E1', title_color=ORANGE,
)
divider()

# ══════════════════════════════════════════════════
#  句型四：職涯轉折正面化句型
# ══════════════════════════════════════════════════
heading1('句型四：職涯轉折正面化句型')

heading2('為什麼空窗期或轉職原因不寫最危險？')
para(
    '你沉默，HR 就替你說話——他們替你說的，通常是：「被裁員了？」「工作出了問題？」「身體有狀況？」這是最糟糕的沉默策略。',
    space_before=4
)
para(
    '很多人在寫履歷時，對於空窗期、頻繁換工作、或「不好解釋」的離職原因，'
    '選擇的策略是：不寫，希望 HR 不要問。'
)
para(
    '當 HR 發現時間軸上有一段空白，他們不會跳過，他們會自己腦補——'
    '通常腦補的都是負面劇情。而他們說的，通常對你不利。'
)

p_key4 = doc.add_paragraph()
p_key4.paragraph_format.space_before = Pt(6)
p_key4.paragraph_format.space_after  = Pt(4)
run = p_key4.add_run(
    '職涯轉折正面化句型的核心邏輯是：主動說明，掌握敘事權，並把「空白」轉化為「選擇」。'
)
set_font(run, 11, bold=True, color=DARK)

heading2('句型公式')
para('用於空窗期：', bold=True, size=10.5, space_before=4, space_after=2)
formula_box('【明確說明空窗起因】＋【期間的主動作為】＋【與目標職位的連結】')
para('用於短暫工作（< 1年）：', bold=True, size=10.5, space_before=8, space_after=2)
formula_box('【加入原因/成果貢獻】＋【離開原因（中性說明）】＋【銜接下一步的邏輯】')

heading3('重點說明：')
bullet('主動說明比被動等問更有力——你掌握了框架')
bullet('「選擇離開」比「被動離職」給讀者完全不同的感受')
bullet('空窗期的「主動作為」不需要很厲害，只要有邏輯的連結即可')

add_hr_visual()
heading2('改寫對照')
add_comparison_table_situation([
    (1, '全職照顧家人後復出',
        '（履歷上直接跳過 2021–2022）',
        '2021–2022 年因擔任主要照顧者暫離全職工作；照顧重心逐步轉移後，2022 年上半年完成 PMP 認證（2022.06）及線上數據分析課程，目前全職求職中，目標為回歸專案管理領域'),
    (2, '任職不到一年離職',
        '（直接列公司名稱與職稱，不加說明）',
        '加入 XX 公司參與新創初期產品建立，完成電商平台 MVP 上線（3 個月），因公司業務方向調整（B2B 轉 B2C）與個人長期目標不符，選擇尋找更匹配的機會'),
    (3, '創業後回求職市場',
        '曾創業，目前尋找機會',
        '2020–2023 年主導個人品牌顧問工作室運營，累計服務 40+ 位求職者，因市場競爭激烈決定加入成熟團隊，將個人學習系統化轉化為企業可規模化的人才培育方案'),
    (4, '頻繁換工作（平均 1–2 年一次）',
        '（未說明原因，直接列公司清單）',
        '過去三份工作均因公司縮編或部門重整離職，每次都在原公司確認結束前完成了主要交付（研發專案進入量產、業務轉型計畫啟動）。目前在找一個能做長一點的位置，目標是在同一個產品或市場方向上真正深下去。'),
])

heading2('練習格式')
add_practice_box([
    '我的「不好解釋」之處是：',
    '□ 空窗期（________年________月 至 ________年________月）',
    '□ 短暫工作（不到 __ 年）',
    '□ 頻繁換工作（平均每 __ 年一次）',
    '',
    '客觀事實是：__________________________________________',
    '',
    '期間我主動做了什麼：___________________________________',
    '',
    '與目標職位的連結是：___________________________________',
    '',
    '正面化後的說明句子：',
    '【中性說明起因】＋【期間作為】＋【連結下一步】',
])

add_tip_box(
    '延伸閱讀',
    [
        '如果你的工作經歷跨越多個產業或職能，可參考搭配指南《跨產業求職：如何讓履歷跨越行業門檻》（獨立文件），',
        '了解如何在整份履歷的結構安排上呼應本句型的轉換框架。',
    ],
    fill='EBF4FB',
    title_color=ACCENT,
)
divider()

# ══════════════════════════════════════════════════
#  句型五：JD 關鍵詞鏡像句型
# ══════════════════════════════════════════════════
heading1('句型五：JD 關鍵詞鏡像句型')

heading2('為什麼投了幾十封都沒有面試，可能跟你的能力無關？')
para(
    '現代大型企業的第一關篩選，很多已不是 HR 人工審查，'
    '而是 ATS（Applicant Tracking System，應徵者追蹤系統）。',
    space_before=4
)

p_key5 = doc.add_paragraph()
p_key5.paragraph_format.space_before = Pt(4)
p_key5.paragraph_format.space_after  = Pt(4)
run = p_key5.add_run(
    'ATS 的運作邏輯很簡單：比對你的履歷和 JD 的關鍵詞重疊率，'
    '重疊率低於門檻，履歷直接被濾掉，HR 連看都不會看到。'
)
set_font(run, 11, bold=True, color=DARK)

para(
    '即使通過 ATS，HR 人工審查時，也習慣先掃「有沒有我要找的那幾個詞」——'
    '因為時間有限，他們看的是信號，不是文章。'
)

p_key5b = doc.add_paragraph()
p_key5b.paragraph_format.space_before = Pt(6)
p_key5b.paragraph_format.space_after  = Pt(4)
run = p_key5b.add_run('JD 關鍵詞鏡像句型的核心邏輯是：用對方的語言說你自己的成果。')
set_font(run, 11, bold=True, color=DARK)

para('這不是複製貼上，而是在你原本的成果描述裡，有意識地加入對方在乎的詞彙。')

heading2('句型公式')
formula_box('Step 1：從 JD 提取 3–5 個核心關鍵詞')
formula_box('Step 2：找出你過去哪個工作經歷「本質上」對應了這個詞')
formula_box('Step 3：把這個詞自然地融入你的成果句子')

# Fix 7 + Fix 26：句型五 適用情境框（加句型二vs五差異）
add_tip_box(
    '【適用情境】',
    [
        '此句型適合所有求職者，尤其是：投遞中大型企業（更可能使用 ATS 系統自動篩選）、',
        '或投遞後長期無回音的求職者。',
        '使用前提是你的經歷本身已具備對應能力，關鍵詞是在真實成果上做語言調整，',
        '而非憑空塞入——後者反而會在面試時露餡。',
        '',
        '與句型二（跨界翻譯）的差異：句型五聚焦「關鍵詞密度對齊」，句型二聚焦「背景框架翻譯」。',
        '若你的背景與目標職位差異較大，建議先用句型二建立整體可讀性，再用句型五微調用詞密度。',
    ],
    fill='EBF4FB',
    title_color=ACCENT,
)

add_hr_visual()
heading2('改寫對照')
add_comparison_table_jd([
    (1, '「客戶關係維護」\n「跨部門溝通」',
        '舉辦客戶活動，反應良好',
        '跨業務、行銷、客服三部門協作，主導規劃 12 場客戶關係維護活動，客戶滿意度 NPS 從 62 提升至 78'),
    (2, '「數據分析」\n「KPI 追蹤」',
        '每週彙整報表，提交主管',
        '每週彙整 5 大指標數據分析報表，建立 KPI 追蹤儀表板，協助主管提前 2 週識別業績落差並啟動應對方案'),
    (3, '「專案管理」\n「時程規劃」',
        '負責協調各部門進度',
        '運用專案管理方法論協調跨部門任務，制定 3 個月時程規劃甘特圖，確保 2 個主要里程碑均準時交付'),
])

# Fix 27：練習格式改版 → 三步驟 + 自我檢核
heading2('練習格式')
heading3('Step 1：提取 JD 關鍵詞')
add_practice_box([
    '目標 JD 的 3–5 個核心關鍵詞：',
    '1. ___  2. ___  3. ___  4. ___  5. ___',
])
heading3('Step 2：改寫句子')
add_practice_box([
    '改前（用自己的話描述這段工作經歷）：',
    '_____________________________________________________',
    '',
    '哪 2–3 個關鍵詞可以自然加入這個描述？',
    '_____________________________________________________',
    '',
    '改後（加入關鍵詞後的句子）：',
    '_____________________________________________________',
])
heading3('Step 3：自我檢核')
table_chk = doc.add_table(rows=5, cols=2)
table_chk.alignment = WD_TABLE_ALIGNMENT.CENTER
table_chk.style = 'Table Grid'
chk_hdr = table_chk.rows[0].cells
for cell, txt, fill in zip(chk_hdr, ['檢核項目', '確認'],
                            ['2C3E50', '2C3E50']):
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(txt)
    set_font(run, 9.5, bold=True, color=WHITE)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)
chk_items = [
    '關鍵詞是「自然融入」成果句，而不是直接貼在句尾',
    '這段描述說的仍然是我真實做過的事',
    '面試時如果被問，我能完整解釋這段經歷',
    '改寫後句子讀起來仍然流暢，不生硬',
]
for i, item in enumerate(chk_items):
    row = table_chk.rows[i + 1]
    row.cells[0].paragraphs[0].clear()
    run = row.cells[0].paragraphs[0].add_run(item)
    set_font(run, 9.5, color=DARK)
    row.cells[0].paragraphs[0].paragraph_format.line_spacing = Pt(14)
    row.cells[1].paragraphs[0].clear()
    run2 = row.cells[1].paragraphs[0].add_run('□')
    set_font(run2, 11, color=DARK)
    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
shaded_para(
    '🎯 完成五個句型的改寫之後，把整份履歷送進診斷工具，確認 ATS 層面的最後缺口。',
    'E8F4FD', size=11, bold=True, color=ACCENT, space_before=10, space_after=2
)
shaded_para(
    '▶ AI 履歷快速診斷（免費）：https://resume-diagnosis.vercel.app/',
    'E8F4FD', size=10.5, bold=True, color=ACCENT, space_before=0, space_after=8
)
divider()
# ══════════════════════════════════════════════════
#  結語
# ══════════════════════════════════════════════════
heading1('結語：句型只是開始')
para('這 5 個句型，幫你解決的是表達層的問題——你做了這些事，但還沒有讓它被看見。但還有另一半：你投的職位方向對嗎？HR 看完你的履歷，知道你是誰嗎？這些問題，句型改不了。', space_before=6)

# Fix 28：結語後半段改為真實案例語氣版本
para(
    '我服務過的求職者裡，有一部分人改完句型就開始收到面試邀約——'
    '通常是本來就大致投對方向、只差表達沒到位的那些人。但還有一部分人改完之後，'
    '還是沒有動靜。那種情況，問題通常不在句型，而在「這份履歷投的方向本來就不對」'
    '或是「HR 根本不知道你要做什麼」——這是另一個問題，句型解決不了。',
    space_before=6
)
para(
    '如果你是後者，我在《跨產業求職履歷重寫指南》裡有更完整的診斷框架，'
    '或者直接找我聊，往往比自己反覆改更快。'
)
para('歡迎透過以下方式聯絡我：', space_before=6)

# AI 診斷 CTA 最顯眼
shaded_para('▶ 立即試用（免費）：AI 履歷快速診斷', 'E8F4FD',
            size=12, bold=True, color=ACCENT, space_before=10, space_after=2)
shaded_para('https://resume-diagnosis.vercel.app/', 'E8F4FD',
            size=10.5, color=ACCENT, space_before=0, space_after=8)
para('📌 加 LINE 洽談：https://lin.ee/IOX6V66', space_before=0, size=10, color=GRAY)
para('📌 官網預約：https://www.careerssl.com/', space_before=0, size=10, color=GRAY)

divider()
para('職涯停看聽 | 蒲朝棟 Tim', size=10, color=GRAY,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10)
para('CDA 認證職涯顧問・104 職涯引導師', size=10, color=GRAY,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0)
para('服務 300+ 位求職者，專注 3-10 年中階職場人才', size=10, color=GRAY,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0)

# ── 儲存（雙路徑：主檔 + v3 命名版）──────────────
base = r'C:\Users\USER\Desktop\職涯停看聽_網站\電子書\5句型電子書'
out_main = os.path.join(base, '讓履歷被看見的5個關鍵句型_完整版.docx')
out_v5   = os.path.join(base, '讓履歷被看見的5個關鍵句型_完整版_v7.docx')
doc.save(out_main)
doc.save(out_v5)
print(f'saved: {out_main}')
print(f'saved: {out_v5}')
