import sys
# Windows cp950 terminal 無法輸出 emoji，強制改為 UTF-8（IMP-063 追溯修正）
if sys.stdout.encoding and sys.stdout.encoding.lower() in ('cp950', 'big5', 'gbk', 'cp936'):
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Color palette ─────────────────────────────────────────────
ORANGE  = RGBColor(0xC4, 0x62, 0x2D)
DARK    = RGBColor(0x1C, 0x1C, 0x1A)
GRAY    = RGBColor(0x6B, 0x6B, 0x68)
LIGHT   = RGBColor(0xF5, 0xF0, 0xE8)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
RED     = RGBColor(0xDC, 0x26, 0x26)
GREEN   = RGBColor(0x05, 0x96, 0x69)

# ── Helper functions ───────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def cell_border(cell, sides=('top','bottom','left','right'), size=6, color='C4622D'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in sides:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    str(size))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18) if level == 1 else Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size  = Pt(18)
        run.font.color.rgb = ORANGE
    elif level == 2:
        run.font.size  = Pt(13)
        run.font.color.rgb = DARK
    else:
        run.font.size  = Pt(11.5)
        run.font.color.rgb = GRAY
    run.font.name = '微軟正黑體'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    return p

def add_body(doc, text, indent=False, color=None):
    p   = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.font.size  = Pt(10.5)
    run.font.color.rgb = color if color else DARK
    run.font.name  = '微軟正黑體'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    return p

def add_bullet(doc, text, sub=False, bold_part=None):
    p   = doc.add_paragraph(style='List Bullet' if not sub else 'List Bullet 2')
    p.paragraph_format.space_after = Pt(3)
    if bold_part and text.startswith(bold_part):
        run = p.add_run(bold_part)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = '微軟正黑體'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
        rest = text[len(bold_part):]
        if rest:
            r2 = p.add_run(rest)
            r2.font.size = Pt(10.5)
            r2.font.name = '微軟正黑體'
            r2._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = '微軟正黑體'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    return p

def add_step_table(doc, steps):
    """steps = list of (step_label, description, note)"""
    table = doc.add_table(rows=len(steps), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [Cm(3.5), Cm(12.5)]
    for i, row in enumerate(table.rows):
        row.cells[0].width = widths[0]
        row.cells[1].width = widths[1]
        step_label, desc, note = steps[i]
        # Left cell (step number)
        set_cell_bg(row.cells[0], 'C4622D')
        p0 = row.cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(step_label)
        r0.bold = True
        r0.font.color.rgb = WHITE
        r0.font.size  = Pt(10.5)
        r0.font.name  = '微軟正黑體'
        r0._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
        # Right cell (description)
        set_cell_bg(row.cells[1], 'FDFCF9')
        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(desc)
        r1.font.size = Pt(10.5)
        r1.font.name = '微軟正黑體'
        r1._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
        if note:
            p1b = row.cells[1].add_paragraph()
            r1b = p1b.add_run(f'  → {note}')
            r1b.font.size  = Pt(9.5)
            r1b.font.color.rgb = GRAY
            r1b.font.name  = '微軟正黑體'
            r1b._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    doc.add_paragraph()

def add_note_box(doc, text, color_hex='FEF3C7', border_color='F59E0B'):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, color_hex)
    cell_border(cell, color=border_color)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = '微軟正黑體'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    doc.add_paragraph()

def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D8D3C8')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ══════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.paragraph_format.space_before = Pt(60)
cover.paragraph_format.space_after  = Pt(0)
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover.add_run('職涯停看聽')
r.font.size = Pt(28)
r.font.color.rgb = ORANGE
r.bold = True
r.font.name = '微軟正黑體'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(0)
r2 = p2.add_run('KIT 電子報 × 免費鉛磁鐵')
r2.font.size = Pt(18)
r2.font.color.rgb = DARK
r2.bold = True
r2.font.name = '微軟正黑體'
r2._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(0)
r3 = p3.add_run('完整操作手冊')
r3.font.size = Pt(18)
r3.font.color.rgb = DARK
r3.bold = True
r3.font.name = '微軟正黑體'
r3._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

doc.add_paragraph()
p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run('建立日期：2025 年 4 月　|　工具：KIT（ConvertKit）+ Vercel 靜態網站')
r4.font.size = Pt(10)
r4.font.color.rgb = GRAY
r4.font.name = '微軟正黑體'
r4._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  CHAPTER 0 — 整體架構說明
# ══════════════════════════════════════════════════════════════
add_heading(doc, '0　整體架構說明')
add_body(doc, '本系統由三個部分組成，彼此串接：')

table0 = doc.add_table(rows=4, cols=3)
table0.style = 'Table Grid'
headers = ['元件', '工具', '用途']
widths0 = [Cm(3), Cm(4.5), Cm(8.5)]
for j, h in enumerate(headers):
    cell = table0.rows[0].cells[j]
    cell.width = widths0[j]
    set_cell_bg(cell, '1C1C1A')
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(10.5)
    r.font.name = '微軟正黑體'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

rows0 = [
    ('網站前端', 'Vercel 靜態網站', '嵌入 KIT 表單，讓訪客填 Email 訂閱'),
    ('Email 自動化', 'KIT（ConvertKit）', '接收訂閱、自動寄送免費電子書、管理電子報名單'),
    ('電子書檔案', 'Google Drive（公開連結）', '儲存 PDF，KIT incentive email 直接連結到此'),
]
for i, (a, b, c) in enumerate(rows0):
    row = table0.rows[i+1]
    for j, text in enumerate([a, b, c]):
        cell = row.cells[j]
        cell.width = widths0[j]
        set_cell_bg(cell, 'F5F0E8' if i % 2 == 0 else 'FDFCF9')
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        r.font.name = '微軟正黑體'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

doc.add_paragraph()

add_body(doc, '訪客流程：')
add_body(doc, '網站填 Email  →  KIT 自動寄出電子書（Incentive Email）  →  訂閱者加入名單  →  後續收到每週電子報', indent=True, color=ORANGE)
doc.add_paragraph()

add_hr(doc)

# ══════════════════════════════════════════════════════════════
#  CHAPTER 1 — KIT 帳號基本設定
# ══════════════════════════════════════════════════════════════
add_heading(doc, '1　KIT 帳號基本設定')
add_body(doc, '登入網址：https://app.kit.com/dashboard')
add_body(doc, '帳號名稱：職涯停看聽（Tim 的帳號）')
doc.add_paragraph()

add_heading(doc, '1-1　帳號資料確認', level=2)
add_step_table(doc, [
    ('步驟 1', '右上角頭像 → Settings → Profile', '確認帳號名稱與 Email 正確'),
    ('步驟 2', 'Email Sending → From Name 填「Tim｜職涯停看聽」', '讓收件人知道是誰寄的'),
    ('步驟 3', 'Reply-to 填你的實際 Email', '訂閱者可以直接回信給你'),
])

add_heading(doc, '1-2　時區設定', level=2)
add_step_table(doc, [
    ('步驟 1', 'Settings → General → Timezone', ''),
    ('步驟 2', '選 Asia/Taipei（UTC+8）', '電子報排程時間才會正確'),
])

add_hr(doc)

# ══════════════════════════════════════════════════════════════
#  CHAPTER 2 — 免費鉛磁鐵表單（已完成）
# ══════════════════════════════════════════════════════════════
add_heading(doc, '2　免費鉛磁鐵表單設定')
add_body(doc, '目前已完成設定的表單：「面試高分回答框架（精華 3 則）」')
add_body(doc, '表單 ID：9309490　|　UID：72a13629c5')
doc.add_paragraph()

add_heading(doc, '2-1　表單目前設定狀態', level=2)
table_status = doc.add_table(rows=5, cols=3)
table_status.style = 'Table Grid'
for j, h in enumerate(['設定項目', '目前值', '狀態']):
    cell = table_status.rows[0].cells[j]
    set_cell_bg(cell, '1C1C1A')
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(10.5)
    r.font.name = '微軟正黑體'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

status_rows = [
    ('Opt-in 模式', 'Auto-confirm（Single opt-in）', '✅ 已設定'),
    ('Incentive Email', '已開啟，寄出免費 PDF', '✅ 已設定'),
    ('成功訊息', '已送出！請收你的 Email，免費框架馬上到你手上。如果沒看到，記得查一下垃圾信件夾。', '✅ 已設定'),
    ('網站嵌入', '嵌入於「學習資源」區塊，面試框架卡片', '✅ 已完成'),
]
for i, (a, b, c) in enumerate(status_rows):
    row = table_status.rows[i+1]
    set_cell_bg(row.cells[0], 'F5F0E8' if i % 2 == 0 else 'FDFCF9')
    set_cell_bg(row.cells[1], 'F5F0E8' if i % 2 == 0 else 'FDFCF9')
    set_cell_bg(row.cells[2], 'F5F0E8' if i % 2 == 0 else 'FDFCF9')
    for j, text in enumerate([a, b, c]):
        p = row.cells[j].paragraphs[0]
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.name = '微軟正黑體'
        if j == 2:
            r.font.color.rgb = GREEN
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

doc.add_paragraph()

add_heading(doc, '2-2　未來新增免費產品時，新建表單的步驟', level=2)
add_step_table(doc, [
    ('步驟 1', 'KIT 後台 → Grow → Forms → + New Form', ''),
    ('步驟 2', '選 Inline（內嵌式）→ 選版型（Clare 或其他）', ''),
    ('步驟 3', '填表單名稱（用產品名稱命名，方便管理）', '例：「履歷句型清單」'),
    ('步驟 4', 'Settings → General → 填寫成功訊息', '提醒收件者查垃圾信件夾'),
    ('步驟 5', 'Settings → Incentive → 勾選 Auto-confirm new subscribers', '改成 Single opt-in'),
    ('步驟 6', 'Settings → Incentive → Edit Email Contents → 上傳 PDF 或填入 Google Drive 連結', ''),
    ('步驟 7', '取得 Embed Code → 複製 HTML 程式碼', ''),
    ('步驟 8', '告知工程師（或自行）嵌入網站對應卡片', '參考第 5 章'),
])

add_note_box(doc, '⚠️  注意：每個免費產品建議建立獨立表單，方便追蹤各產品的轉換率與訂閱來源。')

add_hr(doc)

# ══════════════════════════════════════════════════════════════
#  CHAPTER 3 — Incentive Email（自動寄送電子書）
# ══════════════════════════════════════════════════════════════
add_heading(doc, '3　Incentive Email 設定（自動寄送電子書）')
add_body(doc, 'Incentive Email 是訂閱者填完 Email 後，KIT 自動寄出的第一封信。目前用於交付免費電子書。')
doc.add_paragraph()

add_heading(doc, '3-1　進入 Incentive Email 編輯', level=2)
add_step_table(doc, [
    ('步驟 1', 'KIT → Grow → Forms → 點開對應表單', ''),
    ('步驟 2', 'Settings → Incentive → Edit Email Contents', ''),
    ('步驟 3', '進入 Email 編輯器', ''),
])

add_heading(doc, '3-2　Incentive Email 建議內容結構', level=2)
add_body(doc, '主旨（Subject）範例：')
add_body(doc, '你的「面試高分回答框架（精華 3 則）」來了 ✅', indent=True, color=ORANGE)
doc.add_paragraph()
add_body(doc, '內文建議結構：')

content_items = [
    '開頭：感謝對方，確認他們收到的是什麼',
    '中間：PDF 下載連結（Google Drive 公開連結）',
    '補充：一句話說明完整版（12 則）的價值，附購買方式（加 LINE）',
    '結尾：簡短自我介紹，說明後續電子報內容',
]
for item in content_items:
    add_bullet(doc, item)

doc.add_paragraph()
add_heading(doc, '3-3　Google Drive PDF 連結設定方式', level=2)
add_step_table(doc, [
    ('步驟 1', 'Google Drive → 上傳 PDF 檔案', ''),
    ('步驟 2', '右鍵 → 共用 → 「知道連結的人」都可以檢視', ''),
    ('步驟 3', '複製連結，貼入 Incentive Email 內文', ''),
    ('步驟 4', '測試：用私密模式開啟連結，確認不需要登入即可下載', ''),
])

add_note_box(doc, '💡  建議：Google Drive 連結有時會要求登入。可改用 Google Drive 的「直接下載連結」格式：\n將 /file/d/{ID}/view 改為 /file/d/{ID}/export?format=pdf 或直接用縮網址服務（如 reurl.cc）包裝。')

add_hr(doc)

# ══════════════════════════════════════════════════════════════
#  CHAPTER 4 — 電子報發送流程
# ══════════════════════════════════════════════════════════════
add_heading(doc, '4　電子報發送流程（每週 1–2 封）')
doc.add_paragraph()

add_heading(doc, '4-1　新建一封電子報', level=2)
add_step_table(doc, [
    ('步驟 1', 'KIT → Send → Broadcasts → + New Broadcast', ''),
    ('步驟 2', '填寫主旨（Subject）', '建議：直接點出本期主題，不要太模糊'),
    ('步驟 3', '選擇發送對象：All Subscribers（或特定 Segment）', ''),
    ('步驟 4', '撰寫內文', '參考 4-2 建議格式'),
    ('步驟 5', '右上角 → Send a test → 寄給自己先預覽', '確認排版與連結正常'),
    ('步驟 6', '確認無誤後 → Schedule or Send Now', ''),
])

add_heading(doc, '4-2　電子報內容建議格式', level=2)
add_body(doc, '每期電子報建議維持固定結構，讓讀者知道每次可以期待什麼：')
doc.add_paragraph()

format_rows = [
    ('開頭（3–5 行）', '本期主題一句話說清楚，勾起讀者想讀下去'),
    ('主體（1 個主題）', '一期只講一件事。可以是職涯案例、觀念拆解、工具分享'),
    ('行動呼籲（1 個）', '一期只放一個 CTA，例如：預約諮詢 / 加 LINE / 回信分享'),
    ('結尾署名', 'Tim｜職涯停看聽  + 一句你的核心理念'),
]
table_f = doc.add_table(rows=len(format_rows)+1, cols=2)
table_f.style = 'Table Grid'
for j, h in enumerate(['區塊', '說明']):
    cell = table_f.rows[0].cells[j]
    set_cell_bg(cell, '1C1C1A')
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(10.5)
    r.font.name = '微軟正黑體'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
for i, (a, b) in enumerate(format_rows):
    row = table_f.rows[i+1]
    set_cell_bg(row.cells[0], 'F5F0E8' if i % 2 == 0 else 'FDFCF9')
    set_cell_bg(row.cells[1], 'F5F0E8' if i % 2 == 0 else 'FDFCF9')
    for j, text in enumerate([a, b]):
        p = row.cells[j].paragraphs[0]
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        r.font.name = '微軟正黑體'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

doc.add_paragraph()

add_heading(doc, '4-3　發送頻率建議', level=2)
add_bullet(doc, '初期（前 3 個月）：每週 1 封，固定在同一天發（例如每週三早上）')
add_bullet(doc, '穩定後：可升至每週 2 封，但內容量要夠，不要為發而發')
add_bullet(doc, '主題庫建議：提前準備 4–6 週的主題草稿，避免臨時沒靈感')
doc.add_paragraph()

add_heading(doc, '4-4　電子報主題來源', level=2)
add_bullet(doc, 'Instagram 留言區常見問題 → 展開成一篇')
add_bullet(doc, '客戶諮詢時高頻出現的困境 → 匿名化後分享')
add_bullet(doc, '方格子文章精華 → 改寫成電子報版本')
add_bullet(doc, 'Podcast 集數重點摘要 → 附上連結引流')
add_bullet(doc, '你自己的職涯觀察、轉折故事')

add_hr(doc)

# ══════════════════════════════════════════════════════════════
#  CHAPTER 5 — 後續追蹤序列（Sequence）
# ══════════════════════════════════════════════════════════════
add_heading(doc, '5　後續追蹤序列（建議建立）')
add_body(doc, '新訂閱者在收到免費電子書後，建議透過自動化序列在 3–7 天內追蹤，提升付費轉換。')
doc.add_paragraph()

add_heading(doc, '5-1　建立 Sequence 步驟', level=2)
add_step_table(doc, [
    ('步驟 1', 'KIT → Automate → Sequences → + New Sequence', ''),
    ('步驟 2', '命名：例如「面試框架 → 完整版推薦序列」', ''),
    ('步驟 3', '新增第一封信：Day 0（立即）→ 感謝 + 確認收到', ''),
    ('步驟 4', '新增第二封信：Day 3 → 分享一個面試技巧（建立信任）', ''),
    ('步驟 5', '新增第三封信：Day 5 → 推薦完整版（NT$199），附購買方式', ''),
    ('步驟 6', '將 Sequence 連結到表單：表單設定 → After confirmation → 加入此 Sequence', ''),
])

add_heading(doc, '5-2　第三封推薦信建議內容', level=2)
add_body(doc, '主旨範例：你用了 3 則框架了嗎？完整 12 則等你來拿')
doc.add_paragraph()
add_body(doc, '內文邏輯：')
add_bullet(doc, '回顧：「你前幾天拿到的 3 則框架，是最常被問到的三題」')
add_bullet(doc, '延伸：「但面試不只這三題，還有薪資談判、反問面試官、情境題…」')
add_bullet(doc, '推薦：「完整 12 則版本，NT$199，加 LINE 就可以購買」')
add_bullet(doc, '無壓力結尾：「不想買也沒關係，後續我還會繼續分享職涯技巧」')

add_note_box(doc, '💡  這三封信的目標不是強迫購買，而是讓訂閱者感受到你的內容品質，自然願意付費升級。')

add_hr(doc)

# ══════════════════════════════════════════════════════════════
#  CHAPTER 6 — 網站嵌入說明
# ══════════════════════════════════════════════════════════════
add_heading(doc, '6　KIT 表單嵌入網站說明')
add_body(doc, '目前已完成嵌入。此章節說明未來新增表單時的操作流程。')
doc.add_paragraph()

add_heading(doc, '6-1　取得 Embed Code', level=2)
add_step_table(doc, [
    ('步驟 1', 'KIT → Grow → Forms → 點開要嵌入的表單', ''),
    ('步驟 2', '右上角 Publish → Embed', ''),
    ('步驟 3', '選 HTML → 複製全部程式碼', ''),
])

add_heading(doc, '6-2　CSP 安全政策更新（vercel.json）', level=2)
add_body(doc, '每次新增來自新網域的外部資源，都需要更新 vercel.json 的 Content-Security-Policy：')
doc.add_paragraph()

csp_table = doc.add_table(rows=3, cols=2)
csp_table.style = 'Table Grid'
for j, h in enumerate(['需要允許的來源', '加入的 CSP 欄位']):
    cell = csp_table.rows[0].cells[j]
    set_cell_bg(cell, '1C1C1A')
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(10.5)
    r.font.name = '微軟正黑體'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
csp_rows = [
    ('https://f.convertkit.com（KIT 腳本）', 'script-src'),
    ('https://app.kit.com（表單提交）', 'connect-src'),
]
for i, (a, b) in enumerate(csp_rows):
    row = csp_table.rows[i+1]
    set_cell_bg(row.cells[0], 'F5F0E8')
    set_cell_bg(row.cells[1], 'F5F0E8')
    for j, text in enumerate([a, b]):
        p = row.cells[j].paragraphs[0]
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.name = '微軟正黑體'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

doc.add_paragraph()
add_note_box(doc, '⚠️  KIT 的兩個網域（f.convertkit.com 和 app.kit.com）已在目前的 vercel.json 中加入，無需重複設定。', color_hex='DCFCE7', border_color='16A34A')

add_heading(doc, '6-3　部署指令', level=2)
add_body(doc, '修改完 index.html 或 vercel.json 後，在終端機執行：')
p_code = doc.add_paragraph()
p_code.paragraph_format.left_indent = Cm(0.8)
r_code = p_code.add_run('npx vercel --prod --scope shoppy09-2874s-projects')
r_code.font.name = 'Courier New'
r_code.font.size = Pt(10)
r_code.font.color.rgb = ORANGE
doc.add_paragraph()

add_hr(doc)

# ══════════════════════════════════════════════════════════════
#  CHAPTER 7 — 訂閱者管理
# ══════════════════════════════════════════════════════════════
add_heading(doc, '7　訂閱者管理')
doc.add_paragraph()

add_heading(doc, '7-1　查看訂閱者名單', level=2)
add_step_table(doc, [
    ('路徑', 'KIT → Subscribers → List', ''),
    ('篩選', '右側可按 Tags、Segments 篩選', ''),
    ('狀態說明', 'Confirmed = 正常訂閱　Unconfirmed = 未確認　Unsubscribed = 已退訂', ''),
])

add_heading(doc, '7-2　Tag 標籤建議', level=2)
add_body(doc, '建議用 Tag 區分訂閱者來源，方便之後精準發送：')
tag_rows = [
    ('lead-free-interview', '從「面試高分回答框架」免費索取進來的'),
    ('lead-free-resume', '從「履歷句型清單」進來的（未來）'),
    ('customer-paid', '曾購買付費產品的'),
    ('newsletter-only', '直接從電子報連結訂閱，未索取免費產品的'),
]
for tag, desc in tag_rows:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f'{tag}')
    r1.font.name = 'Courier New'
    r1.font.size = Pt(10)
    r1.font.color.rgb = ORANGE
    r2 = p.add_run(f'　→ {desc}')
    r2.font.size = Pt(10.5)
    r2.font.name = '微軟正黑體'
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

doc.add_paragraph()

add_heading(doc, '7-3　退訂處理', level=2)
add_body(doc, '訂閱者點擊信中的 Unsubscribe 連結後，KIT 自動處理，你不需要手動操作。KIT 法規上要求每封信都必須有退訂連結，這已內建在系統中。')

add_hr(doc)

# ══════════════════════════════════════════════════════════════
#  CHAPTER 8 — 數據追蹤與優化
# ══════════════════════════════════════════════════════════════
add_heading(doc, '8　數據追蹤與優化')
doc.add_paragraph()

add_heading(doc, '8-1　每月應該看的數字', level=2)
metrics_rows = [
    ('Open Rate（開信率）', '30% 以上算健康', '低於 20% → 考慮優化主旨'),
    ('Click Rate（點擊率）', '3–5% 以上算健康', '低 → 檢查 CTA 是否清楚'),
    ('New Subscribers', '每月新增人數', '觀察哪個管道帶來最多訂閱'),
    ('Unsubscribe Rate', '低於 0.5% 算正常', '偏高 → 重新檢視內容是否符合受眾期待'),
]
table_m = doc.add_table(rows=len(metrics_rows)+1, cols=3)
table_m.style = 'Table Grid'
for j, h in enumerate(['指標', '健康標準', '行動建議']):
    cell = table_m.rows[0].cells[j]
    set_cell_bg(cell, '1C1C1A')
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(10.5)
    r.font.name = '微軟正黑體'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
for i, (a, b, c) in enumerate(metrics_rows):
    row = table_m.rows[i+1]
    bg = 'F5F0E8' if i % 2 == 0 else 'FDFCF9'
    for j, text in enumerate([a, b, c]):
        set_cell_bg(row.cells[j], bg)
        p = row.cells[j].paragraphs[0]
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.name = '微軟正黑體'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

doc.add_paragraph()

add_heading(doc, '8-2　主旨 A/B 測試（進階）', level=2)
add_body(doc, 'KIT 付費方案支援 A/B 測試主旨。免費方案則可以手動交替測試：')
add_bullet(doc, '奇數週用問句主旨：「你最怕被問哪一題？」')
add_bullet(doc, '偶數週用結論主旨：「面試被刷掉的真正原因，不是緊張」')
add_bullet(doc, '3 個月後比較兩種風格的開信率')

add_hr(doc)

# ══════════════════════════════════════════════════════════════
#  CHAPTER 9 — 常見問題排除
# ══════════════════════════════════════════════════════════════
add_heading(doc, '9　常見問題排除')
doc.add_paragraph()

faq_items = [
    ('Q：訂閱者填完 Email 後沒收到信',
     '1. 確認 Incentive → Auto-confirm 已勾選\n2. 確認 Send incentive email 已勾選\n3. 叫對方查垃圾信件夾\n4. KIT 後台 → Subscribers → 找到那個 Email → 確認狀態是 Confirmed'),
    ('Q：後台看不到新訂閱者',
     '確認 Subscribers 頁面的篩選器是「All Subscribers」而不是只看「Confirmed」'),
    ('Q：表單在網站上沒有顯示',
     '1. 確認 vercel.json 的 CSP 已加入 f.convertkit.com\n2. 確認 index.html 的 <head> 有加入 KIT 的 <script> 標籤\n3. 重新部署（執行 vercel --prod 指令）'),
    ('Q：想修改表單的成功訊息',
     'KIT → Forms → 點開表單 → Settings → General → Show a success message → 修改文字 → Save'),
    ('Q：想停止某個免費產品的自動寄送',
     'KIT → Forms → 點開表單 → Settings → Incentive → 取消勾選 Send incentive email → Save'),
    ('Q：想換電子書的 PDF 檔案',
     '更新 Google Drive 上的檔案（保持同一個分享連結），或在 Incentive Email 裡換掉連結'),
]

for q, a in faq_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    rq = p.add_run(q)
    rq.bold = True
    rq.font.size = Pt(10.5)
    rq.font.color.rgb = DARK
    rq.font.name = '微軟正黑體'
    rq._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

    for line in a.split('\n'):
        pa = doc.add_paragraph()
        pa.paragraph_format.left_indent = Cm(0.8)
        pa.paragraph_format.space_after = Pt(2)
        ra = pa.add_run(line)
        ra.font.size = Pt(10)
        ra.font.color.rgb = GRAY
        ra.font.name = '微軟正黑體'
        ra._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

add_hr(doc)

# ══════════════════════════════════════════════════════════════
#  CHAPTER 10 — 快速操作速查表
# ══════════════════════════════════════════════════════════════
add_heading(doc, '10　快速操作速查表')
doc.add_paragraph()

quick_rows = [
    ('寄送本週電子報', 'Send → Broadcasts → + New Broadcast'),
    ('查看訂閱者名單', 'Subscribers → List'),
    ('修改自動電子書信', 'Forms → 點表單 → Settings → Incentive → Edit Email Contents'),
    ('新增免費產品表單', 'Grow → Forms → + New Form'),
    ('查看開信率統計', 'Send → Broadcasts → 點對應的 Broadcast'),
    ('建立自動化序列', 'Automate → Sequences → + New Sequence'),
    ('部署網站更新', '終端機執行：npx vercel --prod --scope shoppy09-2874s-projects'),
    ('KIT 登入網址', 'https://app.kit.com/dashboard'),
    ('網站網址', 'https://tzlth-website.vercel.app'),
]

table_q = doc.add_table(rows=len(quick_rows)+1, cols=2)
table_q.style = 'Table Grid'
for j, h in enumerate(['操作', '路徑 / 指令']):
    cell = table_q.rows[0].cells[j]
    set_cell_bg(cell, 'C4622D')
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(10.5)
    r.font.name = '微軟正黑體'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

for i, (a, b) in enumerate(quick_rows):
    row = table_q.rows[i+1]
    bg = 'F5F0E8' if i % 2 == 0 else 'FDFCF9'
    set_cell_bg(row.cells[0], bg)
    set_cell_bg(row.cells[1], bg)
    for j, text in enumerate([a, b]):
        p = row.cells[j].paragraphs[0]
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.name = '微軟正黑體' if j == 0 else 'Courier New'
        if j == 1:
            r.font.color.rgb = ORANGE
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

doc.add_paragraph()

# ── Save ──────────────────────────────────────────────────────
output_path = r'C:\Users\USER\Desktop\職涯停看聽_KIT電子報操作手冊.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
